"""
Generate Methodology Documentation Word Document
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def set_times_new_roman(run):
    """Set font to Times New Roman"""
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def add_heading(doc, text, level=1):
    """Add a heading with Times New Roman font"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph with Times New Roman font"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_times_new_roman(run)
    run.font.bold = bold
    run.font.italic = italic
    return para

def add_bullet_point(doc, text):
    """Add a bullet point with Times New Roman font"""
    para = doc.add_paragraph(text, style='List Bullet')
    for run in para.runs:
        set_times_new_roman(run)
    return para

def add_numbered_point(doc, text):
    """Add a numbered point with Times New Roman font"""
    para = doc.add_paragraph(text, style='List Number')
    for run in para.runs:
        set_times_new_roman(run)
    return para

def create_methodology_documentation():
    """Create comprehensive methodology documentation"""

    doc = Document()

    # Set default font for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Title
    title = doc.add_heading('Open-Ended Coding Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(20)
        run.font.bold = True

    subtitle = doc.add_paragraph('Methodology Documentation and Technical Handbook')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.italic = True

    doc.add_page_break()

    # Table of Contents
    add_heading(doc, 'Table of Contents', 1)
    add_paragraph(doc, '1. Executive Summary')
    add_paragraph(doc, '2. Methodology Documentation')
    add_paragraph(doc, '   2.1 TF-IDF with K-Means Clustering')
    add_paragraph(doc, '   2.2 Latent Dirichlet Allocation (LDA)')
    add_paragraph(doc, '   2.3 Non-negative Matrix Factorization (NMF)')
    add_paragraph(doc, '   2.4 Traditional Keyword-Based Coding')
    add_paragraph(doc, '3. Developer Technical Reference')
    add_paragraph(doc, '   3.1 Core APIs and Classes')
    add_paragraph(doc, '   3.2 ML Analysis Functions')
    add_paragraph(doc, '   3.3 Parameters and Configuration')
    add_paragraph(doc, '4. User Guidance for Analysts')
    add_paragraph(doc, '   4.1 How to Run Analysis')
    add_paragraph(doc, '   4.2 Interpreting Results')
    add_paragraph(doc, '   4.3 Best Practices')
    add_paragraph(doc, '5. Consolidated Reference')
    add_paragraph(doc, '   5.1 15 Essential Outputs')
    add_paragraph(doc, '   5.2 Quality Metrics')
    add_paragraph(doc, '   5.3 Troubleshooting Guide')

    doc.add_page_break()

    # 1. Executive Summary
    add_heading(doc, '1. Executive Summary', 1)
    add_paragraph(doc, 'The Open-Ended Coding Analysis framework is a comprehensive Python-based solution for analyzing qualitative research data through systematic coding, theme identification, and hierarchical categorization. This framework provides researchers with both traditional keyword-based and machine learning-powered approaches to automatically discover patterns, themes, and codes in open-ended text responses.')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Key Features:', bold=True)
    add_bullet_point(doc, 'Three machine learning algorithms for automatic code discovery')
    add_bullet_point(doc, 'Traditional keyword-based coding for manual control')
    add_bullet_point(doc, 'Generation of 15 essential outputs for complete qualitative analysis')
    add_bullet_point(doc, 'Interactive web interface via Streamlit')
    add_bullet_point(doc, 'Programmatic access via Jupyter notebooks')
    add_bullet_point(doc, 'Publication-ready visualizations and reports')
    add_bullet_point(doc, 'Support for multiple data formats (CSV, Excel, JSON, SQL databases)')

    doc.add_page_break()

    # 2. Methodology Documentation
    add_heading(doc, '2. Methodology Documentation', 1)

    # 2.1 TF-IDF with K-Means
    add_heading(doc, '2.1 TF-IDF with K-Means Clustering', 2)

    add_paragraph(doc, 'Objectives:', bold=True)
    add_paragraph(doc, 'This methodology automatically discovers themes and codes in qualitative data by converting text responses into numerical vectors and grouping similar responses together. It is the default and recommended method for exploratory qualitative analysis.')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Theoretical Foundation:', bold=True)
    add_paragraph(doc, 'TF-IDF (Term Frequency-Inverse Document Frequency) is a statistical measure that evaluates the importance of a word in a document relative to a collection of documents. K-Means clustering then groups responses with similar TF-IDF profiles into distinct themes or codes.')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Key Assumptions:', bold=True)
    add_bullet_point(doc, 'Responses containing similar words represent similar themes')
    add_bullet_point(doc, 'The number of codes can be pre-specified or optimized')
    add_bullet_point(doc, 'Each response can be assigned to one or more codes based on confidence thresholds')
    add_bullet_point(doc, 'Word frequency patterns are meaningful indicators of thematic content')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Step-by-Step Process:', bold=True)
    add_numbered_point(doc, 'Text Preprocessing: Remove stop words, convert to lowercase, tokenize responses')
    add_numbered_point(doc, 'TF-IDF Vectorization: Convert text responses into numerical vectors (max 1000 features, 1-2 word n-grams)')
    add_numbered_point(doc, 'K-Means Clustering: Group responses into K clusters (default K=10)')
    add_numbered_point(doc, 'Code Generation: Extract top words from each cluster to create code labels and definitions')
    add_numbered_point(doc, 'Code Assignment: Assign codes to responses with confidence scores based on distance to cluster centroids')
    add_numbered_point(doc, 'Validation: Calculate silhouette scores and other quality metrics')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Confidence Scoring:', bold=True)
    add_paragraph(doc, 'Confidence scores range from 0 to 1 and are calculated based on the distance of each response to the cluster centroid. A score closer to 1 indicates high confidence, while scores below 0.3 may indicate ambiguous or outlier responses.')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Limitations:', bold=True)
    add_bullet_point(doc, 'Requires pre-specification of number of codes (K) or optimization analysis')
    add_bullet_point(doc, 'Assumes spherical cluster shapes; may not capture complex theme structures')
    add_bullet_point(doc, 'Sensitive to initial random seed; results may vary slightly between runs')
    add_bullet_point(doc, 'Performance degrades with very short responses (less than 5 words)')
    add_bullet_point(doc, 'May struggle with highly heterogeneous datasets containing multiple distinct topics')

    add_paragraph(doc, '')
    add_paragraph(doc, 'When to Use:', bold=True)
    add_bullet_point(doc, 'Exploratory analysis of open-ended survey responses')
    add_bullet_point(doc, 'Large datasets (100+ responses) with moderate diversity')
    add_bullet_point(doc, 'When you need fast, interpretable results')
    add_bullet_point(doc, 'When responses are at least 1-2 sentences long')

    doc.add_page_break()

    # 2.2 LDA
    add_heading(doc, '2.2 Latent Dirichlet Allocation (LDA)', 2)

    add_paragraph(doc, 'Objectives:', bold=True)
    add_paragraph(doc, 'LDA is a probabilistic topic modeling technique that assumes documents are mixtures of topics, and topics are mixtures of words. It discovers hidden thematic structures in text data by identifying word patterns that frequently co-occur.')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Theoretical Foundation:', bold=True)
    add_paragraph(doc, 'LDA uses a generative probabilistic model where each document is represented as a probability distribution over topics, and each topic is represented as a probability distribution over words. The algorithm uses Bayesian inference to estimate these distributions.')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Key Assumptions:', bold=True)
    add_bullet_point(doc, 'Documents are composed of multiple topics (not just one)')
    add_bullet_point(doc, 'Topics are defined by word probability distributions')
    add_bullet_point(doc, 'The order of words in documents does not matter (bag-of-words assumption)')
    add_bullet_point(doc, 'Topic distributions follow a Dirichlet prior')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Step-by-Step Process:', bold=True)
    add_numbered_point(doc, 'Text Preprocessing: Tokenize, remove stop words, create document-term matrix')
    add_numbered_point(doc, 'Model Training: Apply LDA algorithm to discover latent topics')
    add_numbered_point(doc, 'Topic-Word Distribution: Calculate probability of each word belonging to each topic')
    add_numbered_point(doc, 'Document-Topic Distribution: Calculate probability of each document belonging to each topic')
    add_numbered_point(doc, 'Code Assignment: Assign codes based on dominant topics (threshold-based)')
    add_numbered_point(doc, 'Interpretation: Generate human-readable code labels from top topic words')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Confidence Scoring:', bold=True)
    add_paragraph(doc, 'Confidence scores represent the probability that a response belongs to a particular topic. LDA naturally provides probabilistic assignments, making confidence scores more theoretically grounded than distance-based methods.')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Limitations:', bold=True)
    add_bullet_point(doc, 'Computationally more expensive than TF-IDF + K-Means')
    add_bullet_point(doc, 'Requires more data for stable results (200+ responses recommended)')
    add_bullet_point(doc, 'Topics may be less interpretable and require more manual labeling')
    add_bullet_point(doc, 'Sensitive to hyperparameters (alpha, beta) which are difficult to optimize')
    add_bullet_point(doc, 'May produce topics that overlap or are difficult to distinguish')

    add_paragraph(doc, '')
    add_paragraph(doc, 'When to Use:', bold=True)
    add_bullet_point(doc, 'When responses are likely to contain multiple themes')
    add_bullet_point(doc, 'For discovery of subtle or hidden patterns in text')
    add_bullet_point(doc, 'When you have larger datasets (200+ responses)')
    add_bullet_point(doc, 'When probabilistic assignments are preferred over hard clustering')

    doc.add_page_break()

    # 2.3 NMF
    add_heading(doc, '2.3 Non-negative Matrix Factorization (NMF)', 2)

    add_paragraph(doc, 'Objectives:', bold=True)
    add_paragraph(doc, 'NMF decomposes the document-term matrix into two non-negative matrices, producing parts-based representations that often yield more interpretable topics than other methods. It is particularly effective for discovering distinct, non-overlapping themes.')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Theoretical Foundation:', bold=True)
    add_paragraph(doc, 'NMF factorizes a matrix V into two matrices W and H such that V approximates W x H, with the constraint that all values must be non-negative. This constraint leads to additive, parts-based representations where topics are built from combinations of words.')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Key Assumptions:', bold=True)
    add_bullet_point(doc, 'Text data can be represented as additive combinations of topics')
    add_bullet_point(doc, 'Non-negativity constraint produces more interpretable results')
    add_bullet_point(doc, 'Topics should be relatively distinct and non-overlapping')
    add_bullet_point(doc, 'Sparse representations are desirable')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Step-by-Step Process:', bold=True)
    add_numbered_point(doc, 'Text Preprocessing: Create TF-IDF weighted document-term matrix')
    add_numbered_point(doc, 'Matrix Factorization: Decompose matrix into topic-word and document-topic matrices')
    add_numbered_point(doc, 'Optimization: Iteratively minimize reconstruction error')
    add_numbered_point(doc, 'Topic Extraction: Identify top words for each topic/code')
    add_numbered_point(doc, 'Code Assignment: Assign codes based on document-topic weights')
    add_numbered_point(doc, 'Labeling: Generate interpretable code labels from topic words')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Confidence Scoring:', bold=True)
    add_paragraph(doc, 'Confidence scores are based on the strength of association between documents and topics in the factorized matrix. Higher weights indicate stronger topical associations.')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Limitations:', bold=True)
    add_bullet_point(doc, 'May not capture topics that naturally overlap')
    add_bullet_point(doc, 'Sensitive to initialization; results may vary between runs')
    add_bullet_point(doc, 'Requires careful selection of number of topics')
    add_bullet_point(doc, 'Less probabilistically interpretable than LDA')
    add_bullet_point(doc, 'Performance depends on proper TF-IDF weighting')

    add_paragraph(doc, '')
    add_paragraph(doc, 'When to Use:', bold=True)
    add_bullet_point(doc, 'When you expect distinct, well-separated themes')
    add_bullet_point(doc, 'When interpretability is a high priority')
    add_bullet_point(doc, 'For shorter documents or responses')
    add_bullet_point(doc, 'When you want sparse topic representations')

    doc.add_page_break()

    # 2.4 Traditional Keyword-Based
    add_heading(doc, '2.4 Traditional Keyword-Based Coding', 2)

    add_paragraph(doc, 'Objectives:', bold=True)
    add_paragraph(doc, 'This methodology allows researchers to apply predefined code frames with manually specified keywords to identify themes in qualitative data. It provides full control over the coding scheme and is ideal for confirmatory research or when specific codes are known in advance.')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Theoretical Foundation:', bold=True)
    add_paragraph(doc, 'Based on traditional content analysis methodology where researchers define codes a priori based on theory, research questions, or preliminary analysis. Keyword matching identifies responses containing specific terms associated with each code.')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Key Assumptions:', bold=True)
    add_bullet_point(doc, 'Codes can be reliably identified through keyword presence')
    add_bullet_point(doc, 'Keywords are representative of underlying themes')
    add_bullet_point(doc, 'Researchers have sufficient domain knowledge to define relevant codes')
    add_bullet_point(doc, 'Responses containing code keywords reflect that theme')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Step-by-Step Process:', bold=True)
    add_numbered_point(doc, 'Code Frame Definition: Create codes with labels, descriptions, and keyword lists')
    add_numbered_point(doc, 'Hierarchical Structure: Optionally define parent-child code relationships')
    add_numbered_point(doc, 'Keyword Matching: Search responses for code keywords (case-insensitive by default)')
    add_numbered_point(doc, 'Code Assignment: Assign all matching codes to each response')
    add_numbered_point(doc, 'Theme Analysis: Group codes into higher-level themes')
    add_numbered_point(doc, 'Categorization: Apply multi-level categorization based on code patterns')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Confidence Scoring:', bold=True)
    add_paragraph(doc, 'Not applicable for keyword-based coding. All matched codes are assigned with binary presence/absence. Confidence can be inferred from number of matching keywords or frequency of keyword mentions.')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Limitations:', bold=True)
    add_bullet_point(doc, 'Requires extensive manual effort to define codes and keywords')
    add_bullet_point(doc, 'May miss responses that express themes using non-keyword language')
    add_bullet_point(doc, 'Cannot discover unexpected or emergent themes')
    add_bullet_point(doc, 'Keyword ambiguity may cause false positives')
    add_bullet_point(doc, 'Less effective for discovering novel patterns in data')

    add_paragraph(doc, '')
    add_paragraph(doc, 'When to Use:', bold=True)
    add_bullet_point(doc, 'Confirmatory research with predefined hypotheses')
    add_bullet_point(doc, 'When specific codes are mandated by research design')
    add_bullet_point(doc, 'For validation of ML-generated codes')
    add_bullet_point(doc, 'When full control over coding scheme is required')
    add_bullet_point(doc, 'Small datasets where ML methods may be unstable')

    doc.add_page_break()

    # 3. Developer Technical Reference
    add_heading(doc, '3. Developer Technical Reference', 1)

    add_heading(doc, '3.1 Core APIs and Classes', 2)

    add_paragraph(doc, 'DataLoader Class', bold=True)
    add_paragraph(doc, 'The DataLoader class handles data ingestion from multiple sources.')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Import:')
    para = add_paragraph(doc, 'from src.data_loader import DataLoader')
    for run in para.runs:
        run.font.name = 'Courier New'

    add_paragraph(doc, '')
    add_paragraph(doc, 'Key Methods:')
    add_bullet_point(doc, 'load_csv(filepath, **kwargs) - Load data from CSV file')
    add_bullet_point(doc, 'load_excel(filepath, sheet_name=0, **kwargs) - Load data from Excel file')
    add_bullet_point(doc, 'load_json(filepath, lines=False, orient=None, **kwargs) - Load JSON data')
    add_bullet_point(doc, 'load_from_sqlite(db_path, query) - Load from SQLite database')
    add_bullet_point(doc, 'load_from_postgres(connection_string, query) - Load from PostgreSQL')
    add_bullet_point(doc, 'validate_dataframe(df, required_columns=None) - Validate data structure')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Example Usage:')
    para = add_paragraph(doc, 'loader = DataLoader()\ndf = loader.load_csv("data/sample_responses.csv")\nis_valid, msg = loader.validate_dataframe(df, required_columns=["response"])')
    for run in para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(10)

    add_paragraph(doc, '')
    add_paragraph(doc, '')
    add_paragraph(doc, 'MLOpenCoder Class', bold=True)
    add_paragraph(doc, 'The MLOpenCoder class implements machine learning-based automatic coding.')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Import:')
    para = add_paragraph(doc, 'from ml_open_coding_analysis import MLOpenCoder')
    for run in para.runs:
        run.font.name = 'Courier New'

    add_paragraph(doc, '')
    add_paragraph(doc, 'Constructor Parameters:')
    add_bullet_point(doc, 'n_codes (int, default=10): Number of codes to discover')
    add_bullet_point(doc, 'method (str, default="tfidf_kmeans"): Algorithm to use - "tfidf_kmeans", "lda", or "nmf"')
    add_bullet_point(doc, 'min_confidence (float, default=0.3): Minimum confidence threshold for code assignment (0-1)')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Key Methods:')
    add_bullet_point(doc, 'fit(responses, stop_words="english") - Train the model on text responses')
    add_bullet_point(doc, 'get_codebook_df() - Get codebook as pandas DataFrame')
    add_bullet_point(doc, 'get_quality_metrics() - Calculate clustering quality metrics')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Example Usage:')
    para = add_paragraph(doc, 'coder = MLOpenCoder(n_codes=10, method="tfidf_kmeans", min_confidence=0.3)\ncoder.fit(df["response"])\ncodebook = coder.get_codebook_df()\nmetrics = coder.get_quality_metrics()')
    for run in para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(10)

    add_paragraph(doc, '')
    add_paragraph(doc, '')
    add_paragraph(doc, 'OpenCodingResults Class', bold=True)
    add_paragraph(doc, 'The OpenCodingResults class generates all 15 essential outputs from coded data.')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Import:')
    para = add_paragraph(doc, 'from ml_open_coding_analysis import OpenCodingResults')
    for run in para.runs:
        run.font.name = 'Courier New'

    add_paragraph(doc, '')
    add_paragraph(doc, 'Constructor Parameters:')
    add_bullet_point(doc, 'df (DataFrame): Original data with responses')
    add_bullet_point(doc, 'coder (MLOpenCoder): Trained coder object')
    add_bullet_point(doc, 'response_col (str, default="response"): Name of text column')
    add_bullet_point(doc, 'id_col (str, default="response_id"): Name of ID column')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Key Methods (15 Outputs):')
    add_bullet_point(doc, 'get_code_assignments() - Response-level codes with confidence scores')
    add_bullet_point(doc, 'get_codebook() - Basic codebook with code definitions')
    add_bullet_point(doc, 'get_codebook_detailed() - Detailed codebook with examples')
    add_bullet_point(doc, 'get_frequency_table() - Code frequency distribution')
    add_bullet_point(doc, 'get_quality_metrics() - Reliability and confidence metrics')
    add_bullet_point(doc, 'get_binary_matrix() - Binary code presence/absence matrix')
    add_bullet_point(doc, 'get_representative_quotes(top_n=5) - Top examples per code')
    add_bullet_point(doc, 'get_cooccurrence_matrix() - Code co-occurrence patterns')
    add_bullet_point(doc, 'get_cooccurrence_pairs(min_count=2) - Frequently co-occurring codes')
    add_bullet_point(doc, 'get_descriptive_stats() - Comprehensive statistics')
    add_bullet_point(doc, 'get_segmentation_analysis(segment_col) - Code patterns by demographics')
    add_bullet_point(doc, 'get_qa_report(sample_size=10) - Quality assurance report')
    add_bullet_point(doc, 'get_uncoded_responses() - Responses without codes')
    add_bullet_point(doc, 'get_low_confidence_responses(threshold=0.5) - Low-confidence assignments')
    add_bullet_point(doc, 'get_ambiguous_responses(min_codes=3) - Multi-coded responses')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Example Usage:')
    para = add_paragraph(doc, 'results = OpenCodingResults(df, coder, response_col="response")\nassignments = results.get_code_assignments()\nfrequency = results.get_frequency_table()\nquotes = results.get_representative_quotes(top_n=5)')
    for run in para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(10)

    doc.add_page_break()

    add_heading(doc, '3.2 ML Analysis Helper Functions', 2)

    add_paragraph(doc, 'The helpers/analysis.py module provides utility functions for ML analysis workflows.')
    add_paragraph(doc, '')

    add_paragraph(doc, 'validate_dataframe(df, required_columns, min_rows=1)', bold=True)
    add_paragraph(doc, 'Validates that a DataFrame meets minimum requirements for analysis.')
    add_paragraph(doc, 'Returns: (is_valid: bool, error_message: str)')
    add_paragraph(doc, '')

    add_paragraph(doc, 'preprocess_responses(df, text_column, remove_nulls=True, remove_duplicates=False, min_length=5)', bold=True)
    add_paragraph(doc, 'Preprocesses text data by removing nulls, duplicates, and short responses.')
    add_paragraph(doc, 'Returns: Preprocessed DataFrame')
    add_paragraph(doc, '')

    add_paragraph(doc, 'find_optimal_codes(df, text_column, min_codes=3, max_codes=15, method="tfidf_kmeans", stop_words="english")', bold=True)
    add_paragraph(doc, 'Analyzes data to find optimal number of codes using silhouette analysis.')
    add_paragraph(doc, 'Returns: (optimal_n_codes: int, analysis_results: dict)')
    add_paragraph(doc, '')

    add_paragraph(doc, 'run_ml_analysis(df, text_column, n_codes=10, method="tfidf_kmeans", min_confidence=0.3, progress_callback=None)', bold=True)
    add_paragraph(doc, 'Runs complete ML analysis pipeline from data to coded results.')
    add_paragraph(doc, 'Returns: (coder: MLOpenCoder, results_df: DataFrame, metrics: dict)')
    add_paragraph(doc, '')

    add_paragraph(doc, 'Example Workflow:')
    para = add_paragraph(doc, 'from helpers.analysis import validate_dataframe, preprocess_responses, run_ml_analysis\n\n# Validate data\nis_valid, msg = validate_dataframe(df, required_columns=["response"])\n\n# Preprocess\ndf_clean = preprocess_responses(df, "response", min_length=5)\n\n# Run analysis\ncoder, results_df, metrics = run_ml_analysis(df_clean, "response", n_codes=10)')
    for run in para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(10)

    doc.add_page_break()

    add_heading(doc, '3.3 Parameters and Configuration', 2)

    add_paragraph(doc, 'Core ML Parameters', bold=True)
    add_paragraph(doc, '')

    # Create parameter table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Parameter'
    header_cells[1].text = 'Default'
    header_cells[2].text = 'Range/Options'
    header_cells[3].text = 'Description'

    # Format header
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.bold = True
                run.font.size = Pt(11)

    # Parameter rows
    params = [
        ('n_codes', '10', '3-30', 'Number of themes/codes to discover'),
        ('method', 'tfidf_kmeans', 'tfidf_kmeans, lda, nmf', 'ML algorithm to use'),
        ('min_confidence', '0.3', '0.1-1.0', 'Confidence threshold for code assignment'),
        ('max_features', '1000', '100-5000', 'Maximum vocabulary size for vectorization'),
        ('stop_words', 'english', 'Language or custom list', 'Stop words to remove from analysis'),
        ('min_df', '2', '1-10', 'Minimum document frequency for terms'),
        ('max_df', '0.8', '0.5-1.0', 'Maximum document frequency for terms'),
        ('ngram_range', '(1, 2)', '(1,1) to (1,3)', 'N-gram range for feature extraction'),
    ]

    for param_data in params:
        row_cells = table.add_row().cells
        for i, text in enumerate(param_data):
            row_cells[i].text = text
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)

    add_paragraph(doc, '')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Recommended Parameter Settings by Use Case:', bold=True)
    add_paragraph(doc, '')

    add_paragraph(doc, 'Small Dataset (30-100 responses):', bold=True)
    add_bullet_point(doc, 'n_codes: 5-8')
    add_bullet_point(doc, 'method: tfidf_kmeans')
    add_bullet_point(doc, 'min_confidence: 0.2')
    add_bullet_point(doc, 'max_features: 500')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Medium Dataset (100-500 responses):', bold=True)
    add_bullet_point(doc, 'n_codes: 8-12')
    add_bullet_point(doc, 'method: tfidf_kmeans or nmf')
    add_bullet_point(doc, 'min_confidence: 0.3')
    add_bullet_point(doc, 'max_features: 1000')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Large Dataset (500+ responses):', bold=True)
    add_bullet_point(doc, 'n_codes: 10-20')
    add_bullet_point(doc, 'method: lda or tfidf_kmeans')
    add_bullet_point(doc, 'min_confidence: 0.3-0.4')
    add_bullet_point(doc, 'max_features: 1500-2000')

    doc.add_page_break()

    # 4. User Guidance for Analysts
    add_heading(doc, '4. User Guidance for Analysts and Clients', 1)

    add_heading(doc, '4.1 How to Run Analysis', 2)

    add_paragraph(doc, 'The framework provides three methods for running analysis, ranging from beginner-friendly to advanced programmatic control.')
    add_paragraph(doc, '')

    add_paragraph(doc, 'Method 1: Streamlit Web Interface (Recommended for Non-Programmers)', bold=True)
    add_paragraph(doc, '')
    add_paragraph(doc, 'Step 1: Launch the Application')
    add_paragraph(doc, 'Open a terminal/command prompt, navigate to the project directory, and run:')
    para = add_paragraph(doc, 'streamlit run app.py')
    for run in para.runs:
        run.font.name = 'Courier New'

    add_paragraph(doc, '')
    add_paragraph(doc, 'Step 2: Upload Your Data')
    add_paragraph(doc, 'The application will open in your web browser. Click "Browse files" to upload a CSV or Excel file containing your open-ended responses. Your file should have at least one column with text responses.')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Step 3: Configure Analysis Parameters')
    add_bullet_point(doc, 'Select the text column containing responses')
    add_bullet_point(doc, 'Choose number of codes (default: 10)')
    add_bullet_point(doc, 'Select analysis method (TF-IDF + K-Means recommended)')
    add_bullet_point(doc, 'Set confidence threshold (default: 0.3)')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Step 4: Run Analysis')
    add_paragraph(doc, 'Click "Start Analysis" and wait for processing to complete. Progress will be displayed on screen.')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Step 5: Review Results')
    add_paragraph(doc, 'The interface will display interactive visualizations, code definitions, frequency tables, and quality metrics.')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Step 6: Export Results')
    add_paragraph(doc, 'Download the complete results package as an Excel file containing all 15 outputs across multiple sheets.')

    add_paragraph(doc, '')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Method 2: Jupyter Notebook (For Researchers and Analysts)', bold=True)
    add_paragraph(doc, '')
    add_paragraph(doc, 'Step 1: Launch Jupyter')
    para = add_paragraph(doc, 'jupyter notebook ml_open_coding_analysis.ipynb')
    for run in para.runs:
        run.font.name = 'Courier New'

    add_paragraph(doc, '')
    add_paragraph(doc, 'Step 2: Load Your Data')
    add_paragraph(doc, 'Modify the data loading cell to point to your CSV/Excel file:')
    para = add_paragraph(doc, 'from src.data_loader import DataLoader\nloader = DataLoader()\ndf = loader.load_csv("path/to/your/data.csv")')
    for run in para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(10)

    add_paragraph(doc, '')
    add_paragraph(doc, 'Step 3: Configure Parameters')
    add_paragraph(doc, 'Adjust the analysis parameters in the configuration cell:')
    para = add_paragraph(doc, 'N_CODES = 10\nMETHOD = "tfidf_kmeans"\nMIN_CONFIDENCE = 0.3\nTEXT_COLUMN = "response"')
    for run in para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(10)

    add_paragraph(doc, '')
    add_paragraph(doc, 'Step 4: Run All Cells')
    add_paragraph(doc, 'Execute all cells in order (Cell > Run All). The notebook will generate all 15 outputs with visualizations.')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Step 5: Review and Export')
    add_paragraph(doc, 'Results are automatically saved to the output directory. Review the generated Excel file and visualizations.')

    add_paragraph(doc, '')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Method 3: Python Script (For Developers)', bold=True)
    add_paragraph(doc, 'For integration into automated workflows, use the Python API directly:')
    para = add_paragraph(doc, 'from helpers.analysis import run_ml_analysis\nfrom ml_open_coding_analysis import OpenCodingResults, ResultsExporter\n\n# Load and analyze\ncoder, results_df, metrics = run_ml_analysis(\n    df, "response", n_codes=10, method="tfidf_kmeans"\n)\n\n# Generate outputs\nresults = OpenCodingResults(df, coder)\nassignments = results.get_code_assignments()\n\n# Export\nexporter = ResultsExporter(results)\nexporter.export_excel("results.xlsx")')
    for run in para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(10)

    doc.add_page_break()

    add_heading(doc, '4.2 Interpreting Results', 2)

    add_paragraph(doc, 'Understanding Code Assignments', bold=True)
    add_paragraph(doc, '')
    add_paragraph(doc, 'The code assignments output shows each response with its assigned codes and confidence scores. Each row represents one response with the following information:')
    add_bullet_point(doc, 'Response ID: Unique identifier for the response')
    add_bullet_point(doc, 'Response Text: The original open-ended response')
    add_bullet_point(doc, 'Assigned Codes: One or more code labels separated by semicolons')
    add_bullet_point(doc, 'Confidence Scores: Numerical values (0-1) indicating assignment certainty')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Confidence Score Interpretation:')
    add_bullet_point(doc, '0.7-1.0: High confidence - response clearly belongs to this code')
    add_bullet_point(doc, '0.5-0.7: Moderate confidence - response likely belongs to this code')
    add_bullet_point(doc, '0.3-0.5: Low confidence - response weakly associated with this code')
    add_bullet_point(doc, 'Below 0.3: Very low confidence - may be noise or outlier (typically excluded)')

    add_paragraph(doc, '')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Understanding the Codebook', bold=True)
    add_paragraph(doc, '')
    add_paragraph(doc, 'The codebook defines each discovered code with:')
    add_bullet_point(doc, 'Code ID: Unique identifier (e.g., CODE_001)')
    add_bullet_point(doc, 'Code Label: Human-readable name (auto-generated from top keywords)')
    add_bullet_point(doc, 'Description: Explanation based on top keywords and patterns')
    add_bullet_point(doc, 'Top Keywords: Most important words defining this code')
    add_bullet_point(doc, 'Example Responses: Representative quotes illustrating the code')
    add_bullet_point(doc, 'Frequency: Number of responses assigned to this code')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Note: Auto-generated code labels should be reviewed and refined by researchers based on domain knowledge and research objectives.')

    add_paragraph(doc, '')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Understanding Quality Metrics', bold=True)
    add_paragraph(doc, '')
    add_paragraph(doc, 'Quality metrics help assess the reliability and validity of the coding:')
    add_paragraph(doc, '')
    add_bullet_point(doc, 'Silhouette Score (-1 to 1): Measures how well-separated codes are. Scores above 0.3 are acceptable, above 0.5 are good.')
    add_bullet_point(doc, 'Average Confidence (0-1): Mean confidence across all code assignments. Higher is better (target: >0.5).')
    add_bullet_point(doc, 'Coverage (%): Percentage of responses that received at least one code. Target: >80%.')
    add_bullet_point(doc, 'Davies-Bouldin Index: Lower is better (well-separated codes). Target: <1.5.')
    add_bullet_point(doc, 'Calinski-Harabasz Score: Higher is better (distinct codes). No fixed threshold.')

    add_paragraph(doc, '')
    add_paragraph(doc, 'If quality metrics are poor:')
    add_bullet_point(doc, 'Try adjusting the number of codes')
    add_bullet_point(doc, 'Switch to a different algorithm')
    add_bullet_point(doc, 'Check if your data has sufficient diversity')
    add_bullet_point(doc, 'Consider data preprocessing (remove very short responses)')

    add_paragraph(doc, '')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Understanding Co-Occurrence Patterns', bold=True)
    add_paragraph(doc, '')
    add_paragraph(doc, 'The co-occurrence matrix and heatmap show how often codes appear together in the same response. High co-occurrence suggests:')
    add_bullet_point(doc, 'Related themes that often co-occur in responses')
    add_bullet_point(doc, 'Potential for higher-level theme groupings')
    add_bullet_point(doc, 'Opportunities to explore relationships between codes')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Low or zero co-occurrence suggests distinct, non-overlapping themes.')

    add_paragraph(doc, '')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Understanding Visualizations', bold=True)
    add_paragraph(doc, '')
    add_bullet_point(doc, 'Frequency Bar Chart: Shows distribution of codes. Look for balanced vs. skewed distributions.')
    add_bullet_point(doc, 'Co-Occurrence Heatmap: Darker colors indicate stronger co-occurrence. Diagonal is always maximum.')
    add_bullet_point(doc, 'Network Diagram: Nodes are codes, edges show co-occurrence. Cluster patterns suggest theme groups.')
    add_bullet_point(doc, 'Word Clouds: Larger words are more important to each code. Useful for quick interpretation.')
    add_bullet_point(doc, 'Confidence Distribution: Should be right-skewed (more high-confidence assignments).')

    doc.add_page_break()

    add_heading(doc, '4.3 Best Practices for Analysts', 2)

    add_paragraph(doc, 'Data Preparation', bold=True)
    add_bullet_point(doc, 'Ensure responses are at least 5-10 words for reliable ML analysis')
    add_bullet_point(doc, 'Remove completely blank or nonsensical responses')
    add_bullet_point(doc, 'Keep responses in their original language (analysis works best in English)')
    add_bullet_point(doc, 'Include at least 30 responses minimum, 100+ recommended for stable results')
    add_bullet_point(doc, 'Preserve original response text without heavy editing')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Choosing Number of Codes', bold=True)
    add_bullet_point(doc, 'Start with 10 codes as a baseline')
    add_bullet_point(doc, 'Use the "Find Optimal Codes" feature in Streamlit to identify best number')
    add_bullet_point(doc, 'For small datasets (30-100), use 5-8 codes')
    add_bullet_point(doc, 'For large datasets (500+), use 15-20 codes')
    add_bullet_point(doc, 'Too few codes: overly broad themes, poor differentiation')
    add_bullet_point(doc, 'Too many codes: fragmented themes, difficult interpretation')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Choosing Analysis Method', bold=True)
    add_bullet_point(doc, 'Default: TF-IDF + K-Means - fastest, most interpretable, good for most cases')
    add_bullet_point(doc, 'Use LDA when: responses likely contain multiple themes, larger datasets available')
    add_bullet_point(doc, 'Use NMF when: you want highly distinct codes, shorter responses, interpretability is critical')
    add_bullet_point(doc, 'Run multiple methods and compare results for robustness')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Setting Confidence Threshold', bold=True)
    add_bullet_point(doc, 'Default 0.3 balances coverage and precision')
    add_bullet_point(doc, 'Lower threshold (0.2): More responses coded but more noise')
    add_bullet_point(doc, 'Higher threshold (0.5): Higher quality but some responses may be uncoded')
    add_bullet_point(doc, 'Review low-confidence responses separately for manual coding')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Validating Results', bold=True)
    add_bullet_point(doc, 'Review representative quotes to ensure codes make sense')
    add_bullet_point(doc, 'Check that code labels accurately reflect the underlying content')
    add_bullet_point(doc, 'Manually review a random sample of 10-20 assignments for accuracy')
    add_bullet_point(doc, 'Examine uncoded and low-confidence responses for patterns')
    add_bullet_point(doc, 'Compare ML results with manual coding on a subset if possible')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Refining Code Labels', bold=True)
    add_bullet_point(doc, 'Auto-generated labels are starting points, not final labels')
    add_bullet_point(doc, 'Rename codes to reflect research-specific terminology')
    add_bullet_point(doc, 'Group related codes into higher-level themes')
    add_bullet_point(doc, 'Add contextual descriptions beyond keywords')
    add_bullet_point(doc, 'Align labels with research questions and framework')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Reporting Results', bold=True)
    add_bullet_point(doc, 'Document methodology: algorithm used, number of codes, parameters')
    add_bullet_point(doc, 'Report quality metrics to establish trustworthiness')
    add_bullet_point(doc, 'Include representative quotes for each code')
    add_bullet_point(doc, 'Visualize frequency distributions and co-occurrence patterns')
    add_bullet_point(doc, 'Discuss limitations: uncoded responses, low-confidence assignments')
    add_bullet_point(doc, 'Explain how codes were refined or validated')

    doc.add_page_break()

    # 5. Consolidated Reference
    add_heading(doc, '5. Consolidated Reference and Knowledge Base', 1)

    add_heading(doc, '5.1 The 15 Essential Outputs', 2)

    add_paragraph(doc, 'The framework generates 15 essential outputs that provide comprehensive qualitative analysis results:')
    add_paragraph(doc, '')

    outputs = [
        ('1. Code Assignments', 'Response-level data showing which codes were assigned to each response with confidence scores. Primary output for further analysis.'),
        ('2. Basic Codebook', 'Simple codebook with code IDs, labels, and descriptions. Quick reference for code definitions.'),
        ('3. Detailed Codebook', 'Comprehensive codebook including top keywords, example responses, and frequencies. Full documentation of coding scheme.'),
        ('4. Frequency Table', 'Statistical distribution showing how many responses received each code. Identifies dominant vs. rare themes.'),
        ('5. Quality Metrics', 'Reliability indicators including silhouette scores, confidence averages, and coverage percentages. Assesses analysis validity.'),
        ('6. Binary Matrix', 'Binary presence/absence matrix for statistical analysis (correlations, chi-square tests, etc.).'),
        ('7. Representative Quotes', 'Top example responses for each code, sorted by confidence. Useful for reporting and presentation.'),
        ('8. Co-Occurrence Matrix', 'Numerical matrix showing how often each pair of codes appears together. Identifies related themes.'),
        ('9. Co-Occurrence Pairs', 'List of code pairs that frequently co-occur, with counts. Highlights theme relationships.'),
        ('10. Descriptive Statistics', 'Summary statistics including total responses, codes per response, average confidence, etc.'),
        ('11. Segmentation Analysis', 'Code patterns broken down by demographic or other grouping variables. Identifies subgroup differences.'),
        ('12. QA Report', 'Quality assurance report with sample assignments for validation. Supports reliability checking.'),
        ('13. Uncoded Responses', 'Responses that did not receive any code assignment. Identifies edge cases and outliers.'),
        ('14. Low-Confidence Responses', 'Responses with confidence scores below threshold. Candidates for manual review.'),
        ('15. Ambiguous Responses', 'Responses assigned to many codes simultaneously. May indicate complex or multifaceted responses.'),
    ]

    for output_name, output_desc in outputs:
        add_paragraph(doc, output_name, bold=True)
        add_paragraph(doc, output_desc)
        add_paragraph(doc, '')

    doc.add_page_break()

    add_heading(doc, '5.2 Quality Metrics Reference', 2)

    add_paragraph(doc, 'Silhouette Score', bold=True)
    add_bullet_point(doc, 'Range: -1 to 1')
    add_bullet_point(doc, 'Interpretation: Measures how well-separated codes are')
    add_bullet_point(doc, 'Good: >0.5')
    add_bullet_point(doc, 'Acceptable: 0.3-0.5')
    add_bullet_point(doc, 'Poor: <0.3')
    add_bullet_point(doc, 'Action if poor: Adjust number of codes or try different algorithm')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Average Confidence', bold=True)
    add_bullet_point(doc, 'Range: 0 to 1')
    add_bullet_point(doc, 'Interpretation: Mean confidence across all assignments')
    add_bullet_point(doc, 'Good: >0.6')
    add_bullet_point(doc, 'Acceptable: 0.4-0.6')
    add_bullet_point(doc, 'Poor: <0.4')
    add_bullet_point(doc, 'Action if poor: Lower min_confidence threshold or increase data quality')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Coverage Percentage', bold=True)
    add_bullet_point(doc, 'Range: 0% to 100%')
    add_bullet_point(doc, 'Interpretation: Percentage of responses that received at least one code')
    add_bullet_point(doc, 'Good: >85%')
    add_bullet_point(doc, 'Acceptable: 70-85%')
    add_bullet_point(doc, 'Poor: <70%')
    add_bullet_point(doc, 'Action if poor: Lower confidence threshold, check data quality, try different method')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Davies-Bouldin Index', bold=True)
    add_bullet_point(doc, 'Range: 0 to infinity (lower is better)')
    add_bullet_point(doc, 'Interpretation: Ratio of within-cluster to between-cluster distances')
    add_bullet_point(doc, 'Good: <1.0')
    add_bullet_point(doc, 'Acceptable: 1.0-1.5')
    add_bullet_point(doc, 'Poor: >1.5')
    add_bullet_point(doc, 'Action if poor: Adjust number of codes')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Calinski-Harabasz Score', bold=True)
    add_bullet_point(doc, 'Range: 0 to infinity (higher is better)')
    add_bullet_point(doc, 'Interpretation: Ratio of between-cluster to within-cluster variance')
    add_bullet_point(doc, 'Good: >100 (dataset dependent)')
    add_bullet_point(doc, 'Acceptable: 50-100')
    add_bullet_point(doc, 'Poor: <50')
    add_bullet_point(doc, 'Action if poor: Adjust number of codes or check data heterogeneity')

    doc.add_page_break()

    add_heading(doc, '5.3 Troubleshooting Guide', 2)

    add_paragraph(doc, 'Issue: Poor Quality Metrics (Low Silhouette Score)', bold=True)
    add_paragraph(doc, 'Possible Causes:')
    add_bullet_point(doc, 'Number of codes not optimal for the data')
    add_bullet_point(doc, 'Dataset too small or homogeneous')
    add_bullet_point(doc, 'Responses too short or low-quality')
    add_paragraph(doc, 'Solutions:')
    add_bullet_point(doc, 'Use "Find Optimal Codes" feature to identify better number')
    add_bullet_point(doc, 'Try different algorithms (LDA, NMF)')
    add_bullet_point(doc, 'Remove very short responses (less than 5 words)')
    add_bullet_point(doc, 'Increase dataset size if possible')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Issue: Many Uncoded Responses', bold=True)
    add_paragraph(doc, 'Possible Causes:')
    add_bullet_point(doc, 'Confidence threshold too high')
    add_bullet_point(doc, 'Responses are outliers or off-topic')
    add_bullet_point(doc, 'Number of codes insufficient to capture diversity')
    add_paragraph(doc, 'Solutions:')
    add_bullet_point(doc, 'Lower min_confidence to 0.2 or 0.25')
    add_bullet_point(doc, 'Increase number of codes')
    add_bullet_point(doc, 'Manually review uncoded responses for patterns')
    add_bullet_point(doc, 'Check for data quality issues')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Issue: Code Labels Not Interpretable', bold=True)
    add_paragraph(doc, 'Possible Causes:')
    add_bullet_point(doc, 'Auto-generated labels based purely on keywords')
    add_bullet_point(doc, 'Codes capture noise rather than meaningful themes')
    add_bullet_point(doc, 'Insufficient context in responses')
    add_paragraph(doc, 'Solutions:')
    add_bullet_point(doc, 'Review representative quotes to understand code meaning')
    add_bullet_point(doc, 'Manually rename codes based on domain knowledge')
    add_bullet_point(doc, 'Try different algorithm or number of codes')
    add_bullet_point(doc, 'Consult top keywords and examples together')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Issue: Codes Are Too Similar', bold=True)
    add_paragraph(doc, 'Possible Causes:')
    add_bullet_point(doc, 'Too many codes for dataset diversity')
    add_bullet_point(doc, 'Data is relatively homogeneous')
    add_paragraph(doc, 'Solutions:')
    add_bullet_point(doc, 'Reduce number of codes')
    add_bullet_point(doc, 'Merge similar codes into broader themes')
    add_bullet_point(doc, 'Check co-occurrence matrix for highly correlated codes')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Issue: Analysis Takes Too Long', bold=True)
    add_paragraph(doc, 'Possible Causes:')
    add_bullet_point(doc, 'Large dataset with many features')
    add_bullet_point(doc, 'LDA method on large data')
    add_paragraph(doc, 'Solutions:')
    add_bullet_point(doc, 'Use TF-IDF + K-Means (fastest method)')
    add_bullet_point(doc, 'Reduce max_features parameter (e.g., 500)')
    add_bullet_point(doc, 'Sample data if dataset is very large (>10,000 responses)')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Issue: Results Not Reproducible', bold=True)
    add_paragraph(doc, 'Possible Causes:')
    add_bullet_point(doc, 'Random initialization in clustering algorithms')
    add_paragraph(doc, 'Solutions:')
    add_bullet_point(doc, 'Set random seed for reproducibility (configure in code)')
    add_bullet_point(doc, 'Run analysis multiple times and compare for stability')
    add_bullet_point(doc, 'If results vary significantly, data may need more preprocessing')

    doc.add_page_break()

    add_heading(doc, '5.4 Data Format Requirements', 2)

    add_paragraph(doc, 'Required Data Structure', bold=True)
    add_paragraph(doc, 'Your input data file (CSV or Excel) must contain at minimum:')
    add_bullet_point(doc, 'One column with text responses (any column name)')
    add_bullet_point(doc, 'At least 30 rows of data (100+ recommended)')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Recommended Data Structure', bold=True)
    add_bullet_point(doc, 'response_id: Unique identifier for each response')
    add_bullet_point(doc, 'response: Text column with open-ended responses')
    add_bullet_point(doc, 'respondent_id: Identifier for the respondent (optional)')
    add_bullet_point(doc, 'timestamp: When response was collected (optional)')
    add_bullet_point(doc, 'Demographic columns: Age, gender, location, etc. (optional, for segmentation)')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Data Quality Guidelines', bold=True)
    add_bullet_point(doc, 'Responses should be at least 5-10 words for reliable analysis')
    add_bullet_point(doc, 'Remove completely blank or null responses before upload')
    add_bullet_point(doc, 'Keep original language (English works best)')
    add_bullet_point(doc, 'Avoid heavy editing or paraphrasing of original responses')
    add_bullet_point(doc, 'Ensure consistent encoding (UTF-8 recommended)')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Supported File Formats', bold=True)
    add_bullet_point(doc, 'CSV (.csv) - Recommended, universal compatibility')
    add_bullet_point(doc, 'Excel (.xlsx, .xls) - Multiple sheets supported')
    add_bullet_point(doc, 'JSON (.json) - For programmatic workflows')
    add_bullet_point(doc, 'SQLite database - Via Python API only')
    add_bullet_point(doc, 'PostgreSQL - Via Python API only')

    doc.add_page_break()

    add_heading(doc, '5.5 Quick Reference: Common Workflows', 2)

    add_paragraph(doc, 'Workflow 1: Basic Exploratory Analysis', bold=True)
    add_numbered_point(doc, 'Prepare CSV file with response column')
    add_numbered_point(doc, 'Launch Streamlit: streamlit run app.py')
    add_numbered_point(doc, 'Upload file and select response column')
    add_numbered_point(doc, 'Use default settings (10 codes, TF-IDF + K-Means, 0.3 confidence)')
    add_numbered_point(doc, 'Run analysis and review quality metrics')
    add_numbered_point(doc, 'Examine frequency chart and codebook')
    add_numbered_point(doc, 'Download results as Excel file')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Workflow 2: Optimized Analysis with Parameter Tuning', bold=True)
    add_numbered_point(doc, 'Prepare and upload data via Streamlit')
    add_numbered_point(doc, 'Click "Find Optimal Number of Codes"')
    add_numbered_point(doc, 'Review silhouette analysis results')
    add_numbered_point(doc, 'Set number of codes to recommended optimal value')
    add_numbered_point(doc, 'Run analysis with optimal parameters')
    add_numbered_point(doc, 'Validate results and export')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Workflow 3: Comparative Analysis Across Methods', bold=True)
    add_numbered_point(doc, 'Run analysis with TF-IDF + K-Means method')
    add_numbered_point(doc, 'Export and save results')
    add_numbered_point(doc, 'Re-run analysis with LDA method')
    add_numbered_point(doc, 'Export and save results')
    add_numbered_point(doc, 'Re-run analysis with NMF method')
    add_numbered_point(doc, 'Compare codebooks and quality metrics across methods')
    add_numbered_point(doc, 'Select best-performing method for final analysis')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Workflow 4: Segmented Analysis by Demographics', bold=True)
    add_numbered_point(doc, 'Ensure demographic columns are in dataset (e.g., age_group, gender)')
    add_numbered_point(doc, 'Run complete analysis on full dataset')
    add_numbered_point(doc, 'Access segmentation analysis output')
    add_numbered_point(doc, 'Specify demographic column for segmentation')
    add_numbered_point(doc, 'Review code patterns across segments')
    add_numbered_point(doc, 'Identify segment-specific themes and differences')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Workflow 5: Manual Validation and Refinement', bold=True)
    add_numbered_point(doc, 'Run initial ML analysis')
    add_numbered_point(doc, 'Export code assignments and codebook')
    add_numbered_point(doc, 'Review representative quotes for each code')
    add_numbered_point(doc, 'Manually check low-confidence responses')
    add_numbered_point(doc, 'Refine code labels based on domain expertise')
    add_numbered_point(doc, 'Merge similar codes or split broad codes as needed')
    add_numbered_point(doc, 'Document final codebook with refined labels')

    doc.add_page_break()

    add_heading(doc, '5.6 Citation and Attribution', 2)

    add_paragraph(doc, 'When using this framework in research publications, please cite appropriately and document your methodology clearly.')
    add_paragraph(doc, '')

    add_paragraph(doc, 'Suggested Methodology Statement:', bold=True)
    add_paragraph(doc, '"Qualitative data were analyzed using machine learning-based open coding implemented via the Open-Ended Coding Analysis framework. [METHOD NAME] was used to automatically discover [N] codes from [N_RESPONSES] open-ended responses. Codes were assigned to responses with a minimum confidence threshold of [THRESHOLD]. The analysis achieved a silhouette score of [SCORE], with [COVERAGE]% of responses receiving at least one code assignment. Code labels were reviewed and refined by researchers to ensure interpretability and alignment with research objectives."')
    add_paragraph(doc, '')

    add_paragraph(doc, 'Replace bracketed placeholders with your specific parameters and results.')

    doc.add_page_break()

    # Footer
    add_heading(doc, 'Document Version Information', 2)
    add_paragraph(doc, 'Document Title: Open-Ended Coding Analysis - Methodology Documentation and Technical Handbook')
    add_paragraph(doc, 'Version: 1.0')
    add_paragraph(doc, 'Date: December 2025')
    add_paragraph(doc, 'Author: JC-OE-Coding Project Team')
    add_paragraph(doc, '')
    add_paragraph(doc, 'This document provides comprehensive methodology documentation, technical reference, and user guidance for the Open-Ended Coding Analysis framework. It is intended for researchers, analysts, developers, and clients using this tool for qualitative data analysis.')

    # Save document
    output_path = '/home/user/JC-OE-Coding/documentation/Methodology_Documentation.docx'
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")
    return output_path

if __name__ == "__main__":
    create_methodology_documentation()
