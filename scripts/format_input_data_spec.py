#!/usr/bin/env python3
"""
Apply comprehensive formatting specifications to Input_Data_Specification.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """
    Set cell border
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    # Create borders element
    tcBorders = OxmlElement('w:tcBorders')

    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            edge_element = OxmlElement(f'w:{edge}')
            edge_element.set(qn('w:val'), edge_data.get('val', 'single'))
            edge_element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            edge_element.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(edge_element)

    tcPr.append(tcBorders)

def set_cell_margins(cell, **margins):
    """
    Set cell margins
    Top, bottom, left, right in DXA (twentieths of a point)
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for margin_name, value in margins.items():
        if value is not None:
            margin_element = OxmlElement(f'w:{margin_name}')
            margin_element.set(qn('w:w'), str(value))
            margin_element.set(qn('w:type'), 'dxa')
            tcMar.append(margin_element)

    tcPr.append(tcMar)

def format_document(doc_path):
    """
    Apply all formatting specifications to the document
    """
    print(f"Opening document: {doc_path}")
    doc = Document(doc_path)

    # Set document margins (0.75 inches = 1080 DXA = 914400 EMU)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Statistics
    stats = {
        'paragraphs': 0,
        'tables': 0,
        'title': 0,
        'heading1': 0,
        'heading2': 0,
        'heading3': 0,
        'body_text': 0,
        'lists': 0
    }

    # Process all paragraphs
    for i, paragraph in enumerate(doc.paragraphs):
        stats['paragraphs'] += 1

        # Determine paragraph type and apply formatting
        style_name = paragraph.style.name

        # Default font settings for all text
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)

        # Title formatting
        if style_name == 'Title' or (i == 0 and paragraph.text.strip()):
            stats['title'] += 1
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(12)  # 240pt in twentieths = 12pt
            paragraph.paragraph_format.space_after = Pt(12)  # 240pt in twentieths = 12pt

            for run in paragraph.runs:
                run.font.size = Pt(28)
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0, 0, 0)

        # Heading 1 formatting
        elif 'Heading 1' in style_name:
            stats['heading1'] += 1
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(12)  # 240pt
            paragraph.paragraph_format.space_after = Pt(6)   # 120pt
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            for run in paragraph.runs:
                run.font.size = Pt(16)
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0, 0, 0)

        # Heading 2 formatting
        elif 'Heading 2' in style_name:
            stats['heading2'] += 1
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(9)   # 180pt
            paragraph.paragraph_format.space_after = Pt(5)    # 100pt
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            for run in paragraph.runs:
                run.font.size = Pt(14)
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0, 0, 0)

        # Heading 3 formatting
        elif 'Heading 3' in style_name:
            stats['heading3'] += 1
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(7)   # 140pt
            paragraph.paragraph_format.space_after = Pt(4)    # 80pt
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            for run in paragraph.runs:
                run.font.size = Pt(13)
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0, 0, 0)

        # TOC formatting
        elif 'TOC' in style_name or 'toc' in style_name.lower():
            paragraph.paragraph_format.line_spacing = 1.5

            for run in paragraph.runs:
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0, 0, 0)

        # Body text / Normal / List formatting
        else:
            stats['body_text'] += 1
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.line_spacing = 1.5

            # Check if it's a list item
            if paragraph.style.name in ['List Bullet', 'List Number', 'List Paragraph']:
                stats['lists'] += 1
                paragraph.paragraph_format.left_indent = Inches(0.5)  # 720pt
                paragraph.paragraph_format.first_line_indent = Inches(-0.25)  # -360pt hanging

            for run in paragraph.runs:
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0, 0, 0)

    # Process all tables
    for table in doc.tables:
        stats['tables'] += 1

        # Process each row
        for row_idx, row in enumerate(table.rows):
            is_header = (row_idx == 0)

            for cell in row.cells:
                # Set cell borders (1pt = 4 in sz units)
                border_config = {
                    'top': {'val': 'single', 'sz': 4, 'color': '000000'},
                    'bottom': {'val': 'single', 'sz': 4, 'color': '000000'},
                    'left': {'val': 'single', 'sz': 4, 'color': '000000'},
                    'right': {'val': 'single', 'sz': 4, 'color': '000000'}
                }
                set_cell_border(cell, **border_config)

                # Set cell margins (in DXA - twentieths of a point)
                set_cell_margins(cell, top=100, bottom=100, left=180, right=180)

                # Format cell background
                if is_header:
                    # Header row: light grey background
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'DDDDDD')
                    cell._element.get_or_add_tcPr().append(shading_elm)
                else:
                    # Data rows: white background
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'FFFFFF')
                    cell._element.get_or_add_tcPr().append(shading_elm)

                # Format cell paragraphs
                for paragraph in cell.paragraphs:
                    # Set line spacing to 1.0 for tables
                    paragraph.paragraph_format.line_spacing = 1.0

                    # Set alignment
                    if is_header:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    # Format runs
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(0, 0, 0)

                        if is_header:
                            run.font.bold = True
                        # For data rows, preserve existing bold/italic but don't force

    # Save the document
    print(f"Saving formatted document to: {doc_path}")
    doc.save(doc_path)

    return stats

if __name__ == '__main__':
    doc_path = '/home/user/JC-OE-Coding/documentation/Input_Data_Specification.docx'

    try:
        stats = format_document(doc_path)

        print("\n" + "="*60)
        print("FORMATTING COMPLETE - SUMMARY")
        print("="*60)
        print(f"Total Paragraphs Processed: {stats['paragraphs']}")
        print(f"  - Title: {stats['title']}")
        print(f"  - Heading 1: {stats['heading1']}")
        print(f"  - Heading 2: {stats['heading2']}")
        print(f"  - Heading 3: {stats['heading3']}")
        print(f"  - Body Text: {stats['body_text']}")
        print(f"  - List Items: {stats['lists']}")
        print(f"Total Tables Formatted: {stats['tables']}")
        print("\nFormatting Applied:")
        print("  - Font: Times New Roman (12pt body, sized headings)")
        print("  - Margins: 0.75 inches on all sides")
        print("  - Line Spacing: 1.5 for body, 1.0 for tables")
        print("  - Table formatting: Borders, headers, cell margins")
        print("  - Heading spacing and alignment")
        print("  - Color scheme: Black text, grey table headers")
        print("="*60)

    except Exception as e:
        print(f"ERROR: {e}")
        import traceback
        traceback.print_exc()
