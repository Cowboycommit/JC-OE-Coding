"""
Apply comprehensive formatting to Benchmark_Standards.docx document.

This script applies all specified formatting including fonts, spacing, margins,
table formatting, and other styling according to the provided specifications.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_cell_border(cell, **kwargs):
    """
    Set cell borders.

    Args:
        cell: Table cell
        kwargs: Border properties (top, bottom, left, right, insideH, insideV)
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    # Create borders element
    tcBorders = OxmlElement('w:tcBorders')

    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), edge_data.get('val', 'single'))
            edge_el.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            edge_el.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(edge_el)

    tcPr.append(tcBorders)


def set_cell_margins(cell, **kwargs):
    """
    Set cell margins.

    Args:
        cell: Table cell
        kwargs: Margin properties (top, bottom, left, right) in twentieths of a point
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for margin in ('top', 'left', 'bottom', 'right'):
        if margin in kwargs:
            mar_el = OxmlElement(f'w:{margin}')
            mar_el.set(qn('w:w'), str(kwargs[margin]))
            mar_el.set(qn('w:type'), 'dxa')
            tcMar.append(mar_el)

    tcPr.append(tcMar)


def set_cell_background(cell, color):
    """
    Set cell background color.

    Args:
        cell: Table cell
        color: Hex color string (e.g., 'DDDDDD')
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    tcPr.append(shading)


def format_document(doc_path):
    """
    Apply comprehensive formatting to the document.

    Args:
        doc_path: Path to the Word document

    Returns:
        Dictionary with formatting statistics
    """
    # Open document
    doc = Document(doc_path)

    # Statistics
    stats = {
        'paragraphs': 0,
        'tables': 0,
        'title_count': 0,
        'heading1_count': 0,
        'heading2_count': 0,
        'heading3_count': 0,
        'body_paragraphs': 0,
        'list_items': 0,
        'toc_items': 0
    }

    # Set document margins (0.75 inches = 1080 DXA)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)

    # Set paragraph spacing for Normal style
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph_format.line_spacing = Pt(18)  # 1.5 * 12pt = 18pt
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Process each paragraph
    for idx, para in enumerate(doc.paragraphs):
        stats['paragraphs'] += 1

        # Determine paragraph type
        is_title = False
        is_heading = False
        heading_level = 0
        is_toc = False
        is_list = False

        # Check if it's a title (first non-empty paragraph with text)
        if idx == 0 and para.text.strip():
            is_title = True
            stats['title_count'] += 1

        # Check if it's a heading based on style
        elif para.style.name.startswith('Heading'):
            is_heading = True
            if 'Heading 1' in para.style.name:
                heading_level = 1
                stats['heading1_count'] += 1
            elif 'Heading 2' in para.style.name:
                heading_level = 2
                stats['heading2_count'] += 1
            elif 'Heading 3' in para.style.name:
                heading_level = 3
                stats['heading3_count'] += 1

        # Check if it's TOC
        elif 'TOC' in para.style.name or any('TOC' in str(run.element.xml) for run in para.runs):
            is_toc = True
            stats['toc_items'] += 1

        # Check if it's a list
        elif para.style.name.startswith('List'):
            is_list = True
            stats['list_items'] += 1

        # Apply formatting based on paragraph type
        if is_title:
            # Title formatting
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(12)  # 240/20 = 12pt
            para.paragraph_format.space_after = Pt(12)   # 240/20 = 12pt
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            para.paragraph_format.line_spacing = Pt(18)

            # If no runs exist, create one with the paragraph text
            if not para.runs:
                para.add_run(para.text)
                para.text = ''  # Clear original text

            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(28)
                run.font.color.rgb = RGBColor(0, 0, 0)
                # Preserve bold/italic if already set

        elif is_heading:
            # Heading formatting
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            para.paragraph_format.line_spacing = Pt(18)

            # If no runs exist, create one with the paragraph text
            if not para.runs:
                para.add_run(para.text)
                para.text = ''  # Clear original text

            if heading_level == 1:
                para.paragraph_format.space_before = Pt(12)  # 240/20 = 12pt
                para.paragraph_format.space_after = Pt(6)     # 120/20 = 6pt
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(16)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.bold = True

            elif heading_level == 2:
                para.paragraph_format.space_before = Pt(9)   # 180/20 = 9pt
                para.paragraph_format.space_after = Pt(5)     # 100/20 = 5pt
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.bold = True

            elif heading_level == 3:
                para.paragraph_format.space_before = Pt(7)   # 140/20 = 7pt
                para.paragraph_format.space_after = Pt(4)     # 80/20 = 4pt
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(13)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.bold = True

        elif is_toc:
            # TOC formatting
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            para.paragraph_format.line_spacing = Pt(18)  # 1.5 spacing

            # Determine TOC level based on indentation or style
            if 'toc 1' in para.style.name.lower() or para.paragraph_format.left_indent == Inches(0):
                para.paragraph_format.left_indent = Inches(0.25)  # 360 DXA
                para.paragraph_format.first_line_indent = Inches(-0.25)  # -360 DXA hanging
            else:
                para.paragraph_format.left_indent = Inches(0.75)  # 1080 DXA
                para.paragraph_format.first_line_indent = Inches(-0.5)  # -720 DXA hanging

            # If no runs exist, create one with the paragraph text
            if not para.runs:
                para.add_run(para.text)
                para.text = ''  # Clear original text

            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)

        elif is_list:
            # List formatting
            para.paragraph_format.left_indent = Inches(0.5)  # 720 DXA
            para.paragraph_format.first_line_indent = Inches(-0.25)  # -360 DXA hanging
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            para.paragraph_format.line_spacing = Pt(18)  # Inherit from body (1.5)

            # If no runs exist, create one with the paragraph text
            if not para.runs:
                para.add_run(para.text)
                para.text = ''  # Clear original text

            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)

        else:
            # Body text formatting
            stats['body_paragraphs'] += 1
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            para.paragraph_format.line_spacing = Pt(18)  # 1.5 * 12pt = 18pt

            # If no runs exist, create one with the paragraph text
            if not para.runs and para.text.strip():
                para.add_run(para.text)
                para.text = ''  # Clear original text

            for run in para.runs:
                # Check if it's code (Courier New)
                if run.font.name and 'Courier' in run.font.name:
                    run.font.name = 'Courier New'
                else:
                    run.font.name = 'Times New Roman'

                if not run.font.size or run.font.size < Pt(10):
                    run.font.size = Pt(12)

                run.font.color.rgb = RGBColor(0, 0, 0)
                # Preserve bold/italic if already set

    # Process tables
    for table_idx, table in enumerate(doc.tables):
        stats['tables'] += 1

        # Add spacing before table (blank paragraph)
        # Note: This is tricky to do after the fact, but we can adjust table spacing

        # Set table properties
        for row_idx, row in enumerate(table.rows):
            is_header = row_idx == 0

            for cell in row.cells:
                # Set cell margins (top: 100pt, bottom: 100pt, left: 180pt, right: 180pt)
                set_cell_margins(cell, top=100, bottom=100, left=180, right=180)

                # Set cell borders (single black line, 1pt)
                border_spec = {'val': 'single', 'sz': 4, 'color': '000000'}
                set_cell_border(cell,
                              top=border_spec,
                              bottom=border_spec,
                              left=border_spec,
                              right=border_spec)

                # Set background color
                if is_header:
                    set_cell_background(cell, 'DDDDDD')
                else:
                    set_cell_background(cell, 'FFFFFF')

                # Format cell paragraphs
                for para in cell.paragraphs:
                    # Set alignment
                    if is_header:
                        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    # Set line spacing (1.0 for tables)
                    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    para.paragraph_format.line_spacing = Pt(12)  # 1.0 * 12pt

                    # Format runs
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(0, 0, 0)

                        if is_header:
                            run.font.bold = True
                        else:
                            # Preserve existing bold/italic for data rows
                            pass

    # Save the document
    doc.save(doc_path)

    return stats


def main():
    """Main function to format the document."""
    doc_path = '/home/user/JC-OE-Coding/documentation/Benchmark_Standards.docx'

    print("Starting document formatting...")
    print(f"Document: {doc_path}")

    if not os.path.exists(doc_path):
        print(f"Error: Document not found at {doc_path}")
        return

    # Apply formatting (this modifies and saves the document)
    stats = format_document(doc_path)

    print("\nFormatting complete!")
    print("\nFormatting Statistics:")
    print(f"  Total Paragraphs: {stats['paragraphs']}")
    print(f"  Title: {stats['title_count']}")
    print(f"  Heading 1: {stats['heading1_count']}")
    print(f"  Heading 2: {stats['heading2_count']}")
    print(f"  Heading 3: {stats['heading3_count']}")
    print(f"  Body Paragraphs: {stats['body_paragraphs']}")
    print(f"  List Items: {stats['list_items']}")
    print(f"  TOC Items: {stats['toc_items']}")
    print(f"  Tables: {stats['tables']}")
    print(f"\nDocument saved to: {doc_path}")


if __name__ == '__main__':
    main()
