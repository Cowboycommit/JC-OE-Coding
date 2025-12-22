"""Shared document creation utilities."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from typing import Optional, List, Dict, Any
from pathlib import Path

from .formatting import set_document_margins, format_table


def create_document(margin_inches: float = 0.75) -> Document:
    """Create a new Word document with default settings.

    Args:
        margin_inches: Document margin size in inches

    Returns:
        New Document instance
    """
    doc = Document()
    set_document_margins(doc, margin_inches)
    return doc


def set_document_defaults(document: Document, font_name: str = "Times New Roman",
                          font_size: int = 12) -> None:
    """Set default font for entire document.

    Args:
        document: Document to modify
        font_name: Default font family
        font_size: Default font size in points
    """
    style = document.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(font_size)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_heading(document: Document, text: str, level: int = 1,
                font_name: str = "Times New Roman",
                font_sizes: Optional[Dict[int, int]] = None,
                space_before: int = 12, space_after: int = 6) -> None:
    """Add a formatted heading to the document.

    Args:
        document: Document to add heading to
        text: Heading text
        level: Heading level (0=Title, 1-3=Heading levels)
        font_name: Font family
        font_sizes: Dict mapping heading levels to font sizes
        space_before: Space before in points
        space_after: Space after in points
    """
    if font_sizes is None:
        font_sizes = {0: 18, 1: 16, 2: 14, 3: 12}

    heading = document.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(space_before)
    heading.paragraph_format.space_after = Pt(space_after)

    for run in heading.runs:
        run.font.name = font_name
        run.font.size = Pt(font_sizes.get(level, 12))
        run.font.color.rgb = RGBColor(0, 0, 0)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_paragraph(document: Document, text: str,
                  font_name: str = "Times New Roman",
                  font_size: int = 12, bold: bool = False,
                  italic: bool = False, alignment: str = "justify",
                  space_before: int = 0, space_after: int = 6,
                  line_spacing: float = 1.5,
                  first_line_indent: Optional[float] = None) -> None:
    """Add a formatted paragraph to the document.

    Args:
        document: Document to add paragraph to
        text: Paragraph text
        font_name: Font family
        font_size: Font size in points
        bold: Whether text is bold
        italic: Whether text is italic
        alignment: Text alignment
        space_before: Space before in points
        space_after: Space after in points
        line_spacing: Line spacing multiplier
        first_line_indent: First line indent in inches
    """
    para = document.add_paragraph()

    # Set paragraph formatting
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing

    if first_line_indent is not None:
        pf.first_line_indent = Inches(first_line_indent)

    alignments = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    pf.alignment = alignments.get(alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Add text with formatting
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_bullet_point(document: Document, text: str,
                     font_name: str = "Times New Roman",
                     font_size: int = 12, level: int = 0,
                     space_after: int = 3) -> None:
    """Add a bullet point to the document.

    Args:
        document: Document to add bullet to
        text: Bullet point text
        font_name: Font family
        font_size: Font size in points
        level: Indentation level (0 = first level)
        space_after: Space after in points
    """
    para = document.add_paragraph(style='List Bullet')
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.left_indent = Inches(0.25 + (level * 0.25))

    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_numbered_item(document: Document, text: str,
                      font_name: str = "Times New Roman",
                      font_size: int = 12, level: int = 0,
                      space_after: int = 3) -> None:
    """Add a numbered list item to the document.

    Args:
        document: Document to add item to
        text: Item text
        font_name: Font family
        font_size: Font size in points
        level: Indentation level
        space_after: Space after in points
    """
    para = document.add_paragraph(style='List Number')
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.left_indent = Inches(0.25 + (level * 0.25))

    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_table(document: Document, headers: List[str],
              rows: List[List[str]], header_bg: str = "#1f77b4",
              header_fg: str = "#ffffff", font_name: str = "Times New Roman",
              font_size: int = 11) -> None:
    """Add a formatted table to the document.

    Args:
        document: Document to add table to
        headers: List of header cell texts
        rows: List of rows, each row is a list of cell texts
        header_bg: Header background color (hex)
        header_fg: Header text color (hex)
        font_name: Font family
        font_size: Font size in points
    """
    num_cols = len(headers)
    table = document.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = 'Table Grid'

    # Add headers
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header

    # Add data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                row_cells[col_idx].text = str(cell_text)

    # Apply formatting
    format_table(table, header_bg=header_bg, header_fg=header_fg,
                 font_name=font_name, font_size=font_size)


def add_code_block(document: Document, code: str,
                   font_name: str = "Courier New",
                   font_size: int = 10) -> None:
    """Add a code block to the document.

    Args:
        document: Document to add code to
        code: Code text
        font_name: Monospace font family
        font_size: Font size in points
    """
    para = document.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.left_indent = Inches(0.25)

    run = para.add_run(code)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def save_document(document: Document, path: str) -> None:
    """Save document to file, creating directories if needed.

    Args:
        document: Document to save
        path: Output file path
    """
    output_path = Path(path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    print(f"Document saved to: {output_path}")
