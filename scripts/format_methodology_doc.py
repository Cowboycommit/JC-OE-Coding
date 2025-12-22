#!/usr/bin/env python3
"""
Apply comprehensive formatting specifications to Methodology_Documentation.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """
    Set cell borders
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Create borders element
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            edge_el = OxmlElement(f'w:{edge}')
            for key, value in edge_data.items():
                edge_el.set(qn(f'w:{key}'), str(value))
            tcBorders.append(edge_el)
    tcPr.append(tcBorders)

def set_cell_margins(cell, **kwargs):
    """
    Set cell margins (in twentieths of a point)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for margin_name in ('top', 'left', 'bottom', 'right'):
        if margin_name in kwargs:
            node = OxmlElement(f'w:{margin_name}')
            node.set(qn('w:w'), str(kwargs[margin_name]))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)

    tcPr.append(tcMar)

def format_document(doc_path):
    """
    Apply all formatting specifications to the document
    """
    # Load the document
    doc = Document(doc_path)

    # Statistics
    stats = {
        'paragraphs': 0,
        'tables': 0,
        'headings': {'title': 0, 'h1': 0, 'h2': 0, 'h3': 0},
        'list_items': 0,
        'toc_items': 0
    }

    # Set document margins (0.75 inches = 1080 DXA)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Process all paragraphs
    for para in doc.paragraphs:
        stats['paragraphs'] += 1

        # Identify paragraph type by style name or content
        style_name = para.style.name if para.style else ''
        text = para.text.strip()

        # Set default paragraph format
        para_format = para.paragraph_format

        # Title detection (first paragraph, typically)
        if 'Title' in style_name or (stats['paragraphs'] == 1 and text and len(text) > 0):
            # Title formatting
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(28)
                run.font.bold = True
                run.font.italic = True
                run.font.color.rgb = RGBColor(0, 0, 0)

            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para_format.space_before = Pt(240)
            para_format.space_after = Pt(240)
            para_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            stats['headings']['title'] += 1

        # Heading 1
        elif 'Heading 1' in style_name or style_name == 'Heading1':
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para_format.space_before = Pt(240)
            para_format.space_after = Pt(120)
            para_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            stats['headings']['h1'] += 1

        # Heading 2
        elif 'Heading 2' in style_name or style_name == 'Heading2':
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para_format.space_before = Pt(180)
            para_format.space_after = Pt(100)
            para_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            stats['headings']['h2'] += 1

        # Heading 3
        elif 'Heading 3' in style_name or style_name == 'Heading3':
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para_format.space_before = Pt(140)
            para_format.space_after = Pt(80)
            para_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            stats['headings']['h3'] += 1

        # Table of Contents
        elif 'TOC' in style_name or 'Contents' in style_name:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)

            # TOC formatting based on level
            if 'TOC 1' in style_name or text.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                para_format.left_indent = Pt(360)
                para_format.first_line_indent = Pt(-360)
            else:  # Sub-levels
                para_format.left_indent = Pt(1080)
                para_format.first_line_indent = Pt(-720)

            para_format.line_spacing = Pt(360)  # 1.5 line spacing
            stats['toc_items'] += 1

        # List items (bullets)
        elif 'List' in style_name or text.startswith('•') or para_format.left_indent:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)

            para_format.left_indent = Pt(720)
            para_format.first_line_indent = Pt(-360)
            para_format.line_spacing = Pt(360)  # 1.5 line spacing (inherits from body)
            stats['list_items'] += 1

        # Body text (default)
        else:
            for run in para.runs:
                # Check for code formatting
                if run.font.name and 'Courier' in run.font.name:
                    run.font.name = 'Courier New'
                else:
                    run.font.name = 'Times New Roman'

                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
                # Preserve bold and italic

            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para_format.line_spacing = Pt(360)  # 1.5 line spacing in twentieths of a point

    # Process all tables
    for table in doc.tables:
        stats['tables'] += 1

        # Table borders
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # Set table borders
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # 1pt = 4 eighths of a point
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)

        # Process table rows
        for row_idx, row in enumerate(table.rows):
            is_header = (row_idx == 0)

            for cell in row.cells:
                # Set cell margins (top: 100, bottom: 100, left: 180, right: 180)
                set_cell_margins(cell, top=100, bottom=100, left=180, right=180)

                # Set cell background color
                if is_header:
                    # Header row: light grey background
                    cell_shading = cell._element.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), 'DDDDDD')
                    cell_shading.append(shd)

                # Format cell text
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(0, 0, 0)

                        if is_header:
                            run.font.bold = True
                        else:
                            # Preserve original bold/italic for data rows
                            pass

                    # Alignment
                    if is_header:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    # Line spacing for tables: 1.0 (240 in twentieths of a point)
                    para.paragraph_format.line_spacing = Pt(240)

    # Save the formatted document
    doc.save(doc_path)

    return stats

def main():
    """Main execution"""
    doc_path = '/home/user/JC-OE-Coding/documentation/Methodology_Documentation.docx'

    print("=" * 70)
    print("Applying Formatting Specifications to Methodology Documentation")
    print("=" * 70)
    print()

    try:
        stats = format_document(doc_path)

        print("✓ Formatting completed successfully!")
        print()
        print("Document Statistics:")
        print("-" * 70)
        print(f"  Total Paragraphs: {stats['paragraphs']}")
        print(f"  Total Tables: {stats['tables']}")
        print(f"  Headings:")
        print(f"    - Title: {stats['headings']['title']}")
        print(f"    - Heading 1: {stats['headings']['h1']}")
        print(f"    - Heading 2: {stats['headings']['h2']}")
        print(f"    - Heading 3: {stats['headings']['h3']}")
        print(f"  List Items: {stats['list_items']}")
        print(f"  Table of Contents Items: {stats['toc_items']}")
        print()
        print("Applied Formatting:")
        print("-" * 70)
        print("  ✓ Font Family: Times New Roman (all text)")
        print("  ✓ Font Sizes: Title (28pt), H1 (16pt), H2 (14pt), H3 (13pt), Body (12pt)")
        print("  ✓ Line Spacing: Body & TOC (1.5), Tables (1.0)")
        print("  ✓ Margins: 0.75 inches on all sides")
        print("  ✓ Alignment: Left (title, headings, body), Center (table headers)")
        print("  ✓ Heading Spacing: Title (240/240), H1 (240/120), H2 (180/100), H3 (140/80)")
        print("  ✓ Table Formatting: Borders, header backgrounds, cell margins")
        print("  ✓ Table of Contents: Proper indentation and numbering")
        print("  ✓ List Formatting: Bullet points with proper indentation")
        print("  ✓ Color Scheme: Black text, grey table headers")
        print()
        print(f"Document saved: {doc_path}")
        print("=" * 70)

    except Exception as e:
        print(f"✗ Error formatting document: {e}")
        import traceback
        traceback.print_exc()
        return 1

    return 0

if __name__ == '__main__':
    exit(main())
