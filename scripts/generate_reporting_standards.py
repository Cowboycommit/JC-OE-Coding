"""
Create Reporting and Visualization Standards documentation.

This script generates a comprehensive Word document with reporting and
visualization standards for the Open-Ended Coding Analysis project.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os


def set_times_new_roman(run):
    """Set font to Times New Roman."""
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)


def add_heading(doc, text, level=1):
    """Add a heading with Times New Roman font."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
    return heading


def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph with Times New Roman font."""
    para = doc.add_paragraph(text)
    for run in para.runs:
        set_times_new_roman(run)
        run.font.bold = bold
        run.font.italic = italic
    return para


def add_bullet(doc, text):
    """Add a bulleted list item."""
    para = doc.add_paragraph(text, style='List Bullet')
    for run in para.runs:
        set_times_new_roman(run)
    return para


def add_numbered(doc, text):
    """Add a numbered list item."""
    para = doc.add_paragraph(text, style='List Number')
    for run in para.runs:
        set_times_new_roman(run)
    return para


def add_table_content(doc, headers, rows):
    """Add a table with content."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.bold = True
                run.font.size = Pt(11)

    # Add rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)

    return table


def create_document():
    """Create the Reporting & Visualization Standards document."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Title page
    title = doc.add_heading('Reporting & Visualization Standards', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(20)
        run.font.bold = True

    subtitle = doc.add_paragraph('Open-Ended Coding Analysis Framework')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.italic = True

    doc.add_paragraph()  # Spacing

    version_para = doc.add_paragraph('Version 1.0')
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in version_para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    doc.add_page_break()

    # ============================================================================
    # 1. INTRODUCTION
    # ============================================================================
    add_heading(doc, '1. Introduction', level=1)

    add_paragraph(doc,
        'This document establishes comprehensive reporting and visualization standards '
        'for the Open-Ended Coding Analysis Framework. These standards ensure consistency, '
        'clarity, and professional quality across all analytical outputs, enabling effective '
        'communication of qualitative research findings to diverse stakeholders.')

    doc.add_paragraph()

    add_heading(doc, '1.1 Purpose and Scope', level=2)
    add_paragraph(doc, 'This document provides:')
    add_bullet(doc, 'Standardized templates for graphical visualizations')
    add_bullet(doc, 'Guidelines for reporting analytical techniques and results')
    add_bullet(doc, 'Frameworks for presenting analysis inferences as textual summaries')
    add_bullet(doc, 'Recommendations for integrating results into platform dashboards and user interfaces')

    doc.add_paragraph()

    add_heading(doc, '1.2 Document Organization', level=2)
    add_paragraph(doc, 'This document is organized into the following sections:')
    add_numbered(doc, 'Graphical Templates and Visualization Guidelines')
    add_numbered(doc, 'Reporting Standards for Analytical Techniques')
    add_numbered(doc, 'Analysis Inferences as Textual Summaries')
    add_numbered(doc, 'Platform Dashboard and UX Integration Recommendations')

    doc.add_page_break()

    # ============================================================================
    # 2. GRAPHICAL TEMPLATES AND VISUALIZATION GUIDELINES
    # ============================================================================
    add_heading(doc, '2. Graphical Templates and Visualization Guidelines', level=1)

    add_paragraph(doc,
        'This section establishes standardized templates and design principles for all '
        'visualizations produced by the framework. Consistency in visual presentation '
        'enhances comprehension and professional credibility.')

    doc.add_paragraph()

    add_heading(doc, '2.1 Color Palette Standards', level=2)
    add_paragraph(doc, 'All visualizations should adhere to the following color scheme:')

    # Color palette table
    color_headers = ['Color Name', 'Hex Code', 'Usage Context']
    color_rows = [
        ['Primary Blue', '#1f77b4', 'Main charts, primary data series, interactive elements'],
        ['Secondary Gray', '#f0f2f6', 'Backgrounds, containers, secondary elements'],
        ['Text Dark', '#262730', 'All text labels, annotations, titles'],
        ['Success Green', '#28a745', 'Positive indicators, high confidence scores'],
        ['Warning Yellow', '#ffc107', 'Moderate confidence scores, attention areas'],
        ['Danger Red', '#dc3545', 'Low confidence scores, warnings, errors'],
        ['White', '#ffffff', 'Chart backgrounds, negative space'],
    ]
    add_table_content(doc, color_headers, color_rows)

    doc.add_paragraph()

    add_heading(doc, '2.2 Typography Standards', level=2)
    add_paragraph(doc, 'Typography guidelines for all visualizations:')
    add_bullet(doc, 'Chart Titles: Times New Roman, 14pt, Bold')
    add_bullet(doc, 'Axis Labels: Times New Roman, 11pt, Regular')
    add_bullet(doc, 'Data Labels: Times New Roman, 10pt, Regular')
    add_bullet(doc, 'Legends: Times New Roman, 10pt, Regular')
    add_bullet(doc, 'Annotations: Times New Roman, 9pt, Italic')

    doc.add_paragraph()

    add_heading(doc, '2.3 Standard Visualization Types', level=2)
    add_paragraph(doc,
        'The framework supports the following standardized visualization types, '
        'each serving specific analytical purposes:')

    doc.add_paragraph()

    add_heading(doc, '2.3.1 Code Frequency Bar Charts', level=3)
    add_paragraph(doc, 'Purpose: Display the distribution and frequency of discovered codes.')
    add_paragraph(doc, 'Specifications:')
    add_bullet(doc, 'Orientation: Vertical bars with horizontal code labels')
    add_bullet(doc, 'Color: Primary Blue (#1f77b4) for all bars')
    add_bullet(doc, 'Data Labels: Display count values above each bar')
    add_bullet(doc, 'Sorting: Sort bars in descending order by frequency')
    add_bullet(doc, 'Axis Labels: Y-axis shows "Number of Responses", X-axis shows "Code Label"')
    add_bullet(doc, 'Grid Lines: Horizontal grid lines for readability')
    add_bullet(doc, 'Maximum Codes Displayed: 15-20 codes (top codes by frequency)')

    doc.add_paragraph()

    add_heading(doc, '2.3.2 Co-occurrence Heatmaps', level=3)
    add_paragraph(doc, 'Purpose: Visualize relationships between codes that appear together.')
    add_paragraph(doc, 'Specifications:')
    add_bullet(doc, 'Layout: Square matrix with codes on both axes')
    add_bullet(doc, 'Color Scale: Sequential gradient from light to dark (e.g., Blues or Viridis)')
    add_bullet(doc, 'Annotations: Display numeric co-occurrence counts in cells')
    add_bullet(doc, 'Diagonal: Highlight or remove diagonal (self-co-occurrence)')
    add_bullet(doc, 'Size: Minimum 600x600 pixels for readability')
    add_bullet(doc, 'Labels: Rotate X-axis labels 45 degrees for better legibility')

    doc.add_paragraph()

    add_heading(doc, '2.3.3 Confidence Score Scatter Plots', level=3)
    add_paragraph(doc, 'Purpose: Examine relationship between code frequency and confidence.')
    add_paragraph(doc, 'Specifications:')
    add_bullet(doc, 'X-axis: Code frequency (count)')
    add_bullet(doc, 'Y-axis: Average confidence score')
    add_bullet(doc, 'Point Size: Proportional to code frequency')
    add_bullet(doc, 'Color: Gradient based on confidence level (Red to Yellow to Green)')
    add_bullet(doc, 'Hover Information: Display code label, count, and exact confidence score')
    add_bullet(doc, 'Reference Lines: Add horizontal line at confidence threshold (e.g., 0.5)')

    doc.add_paragraph()

    add_heading(doc, '2.3.4 Distribution Histograms', level=3)
    add_paragraph(doc, 'Purpose: Show distribution of codes per response or confidence scores.')
    add_paragraph(doc, 'Specifications:')
    add_bullet(doc, 'Bar Color: Primary Blue (#1f77b4)')
    add_bullet(doc, 'Bin Size: Automatic or based on data range (typically 10-30 bins)')
    add_bullet(doc, 'Axis Labels: Clear description of what is being measured')
    add_bullet(doc, 'Statistics Overlay: Include mean and median as vertical reference lines')
    add_bullet(doc, 'Grid: Light horizontal grid lines')

    doc.add_paragraph()

    add_heading(doc, '2.3.5 Network Diagrams', level=3)
    add_paragraph(doc, 'Purpose: Visualize code relationships and thematic structures.')
    add_paragraph(doc, 'Specifications:')
    add_bullet(doc, 'Node Size: Proportional to code frequency')
    add_bullet(doc, 'Edge Width: Proportional to co-occurrence strength')
    add_bullet(doc, 'Layout Algorithm: Force-directed or circular layout')
    add_bullet(doc, 'Node Color: Color-coded by theme or cluster')
    add_bullet(doc, 'Labels: Display code labels near or within nodes')
    add_bullet(doc, 'Interactive Elements: Hover for details, click to highlight connections')

    doc.add_paragraph()

    add_heading(doc, '2.4 Interactivity Standards', level=2)
    add_paragraph(doc, 'For web-based visualizations (Streamlit, dashboards):')
    add_bullet(doc, 'Hover Tooltips: Display detailed information on mouse hover')
    add_bullet(doc, 'Zoom and Pan: Enable for large datasets or complex visualizations')
    add_bullet(doc, 'Click Actions: Provide drill-down or filtering capabilities where appropriate')
    add_bullet(doc, 'Legend Interaction: Allow toggling of data series by clicking legend items')
    add_bullet(doc, 'Export Options: Provide download buttons for PNG, SVG, or PDF formats')

    doc.add_paragraph()

    add_heading(doc, '2.5 Accessibility Guidelines', level=2)
    add_bullet(doc, 'Color Blindness: Ensure color schemes are distinguishable for colorblind users')
    add_bullet(doc, 'Text Contrast: Maintain minimum 4.5:1 contrast ratio for all text')
    add_bullet(doc, 'Alternative Text: Provide descriptive alt text for all charts')
    add_bullet(doc, 'Screen Reader Support: Include semantic HTML and ARIA labels')
    add_bullet(doc, 'Keyboard Navigation: Ensure all interactive elements are keyboard accessible')

    doc.add_page_break()

    # ============================================================================
    # 3. REPORTING STANDARDS FOR ANALYTICAL TECHNIQUES
    # ============================================================================
    add_heading(doc, '3. Reporting Standards for Analytical Techniques', level=1)

    add_paragraph(doc,
        'This section defines how analytical techniques, results, and methodologies '
        'should be documented and reported. Clear documentation ensures reproducibility, '
        'transparency, and effective communication of research findings.')

    doc.add_paragraph()

    add_heading(doc, '3.1 General Reporting Principles', level=2)
    add_bullet(doc, 'Transparency: Clearly document all methods, parameters, and assumptions')
    add_bullet(doc, 'Reproducibility: Provide sufficient detail to enable replication')
    add_bullet(doc, 'Completeness: Report both successful and problematic results')
    add_bullet(doc, 'Clarity: Use plain language accessible to non-technical stakeholders')
    add_bullet(doc, 'Context: Always provide interpretation alongside raw numbers')

    doc.add_paragraph()

    add_heading(doc, '3.2 Method Documentation Standards', level=2)
    add_paragraph(doc, 'For each analytical method employed, report the following:')

    method_headers = ['Component', 'Required Information', 'Example']
    method_rows = [
        ['Algorithm Name', 'Specific technique used', 'TF-IDF + K-Means Clustering'],
        ['Version', 'Software and library versions', 'scikit-learn 1.3.0'],
        ['Parameters', 'All configuration settings', 'n_clusters=10, random_state=42'],
        ['Preprocessing', 'Data cleaning steps', 'Removed null values, lowercased text'],
        ['Validation', 'Quality metrics used', 'Silhouette score, coverage percentage'],
    ]
    add_table_content(doc, method_headers, method_rows)

    doc.add_paragraph()

    add_heading(doc, '3.3 Table Reporting Standards', level=2)
    add_paragraph(doc, 'All tables should follow these formatting guidelines:')

    doc.add_paragraph()

    add_heading(doc, '3.3.1 Code Assignment Tables', level=3)
    add_paragraph(doc, 'Purpose: Display individual response-level code assignments.')
    add_paragraph(doc, 'Required Columns:')
    add_bullet(doc, 'Response ID or Index')
    add_bullet(doc, 'Response Text (truncated if necessary)')
    add_bullet(doc, 'Assigned Codes (comma-separated list)')
    add_bullet(doc, 'Confidence Scores (corresponding to each code)')
    add_bullet(doc, 'Number of Codes Assigned')

    add_paragraph(doc, 'Formatting:')
    add_bullet(doc, 'Font: Times New Roman, 10pt')
    add_bullet(doc, 'Header: Bold, 11pt')
    add_bullet(doc, 'Alignment: Left for text, center for numeric values')
    add_bullet(doc, 'Row Limits: Display top 10-20 rows, with option to view full table')

    doc.add_paragraph()

    add_heading(doc, '3.3.2 Codebook Tables', level=3)
    add_paragraph(doc, 'Purpose: Comprehensive reference of all discovered codes.')
    add_paragraph(doc, 'Required Columns:')
    add_bullet(doc, 'Code ID (e.g., CODE_01, CODE_02)')
    add_bullet(doc, 'Code Label (human-readable name)')
    add_bullet(doc, 'Description (brief explanation of code meaning)')
    add_bullet(doc, 'Keywords (top terms associated with code)')
    add_bullet(doc, 'Frequency Count')
    add_bullet(doc, 'Percentage of Total Responses')
    add_bullet(doc, 'Average Confidence Score')
    add_bullet(doc, 'Representative Examples (2-3 quotes)')

    doc.add_paragraph()

    add_heading(doc, '3.3.3 Frequency Tables', level=3)
    add_paragraph(doc, 'Purpose: Statistical summary of code distributions.')
    add_paragraph(doc, 'Required Columns:')
    add_bullet(doc, 'Code ID and Label')
    add_bullet(doc, 'Count (number of responses)')
    add_bullet(doc, 'Percentage (of total responses)')
    add_bullet(doc, 'Average Confidence Score')

    add_paragraph(doc, 'Formatting:')
    add_bullet(doc, 'Sort: Descending order by count')
    add_bullet(doc, 'Percentages: Display with 1 decimal place (e.g., 23.5%)')
    add_bullet(doc, 'Confidence: Display with 2-3 decimal places (e.g., 0.847)')
    add_bullet(doc, 'Conditional Formatting: Use color gradients for percentages and confidence')

    doc.add_paragraph()

    add_heading(doc, '3.3.4 Quality Metrics Tables', level=3)
    add_paragraph(doc, 'Purpose: Report validation and quality assessment metrics.')
    add_paragraph(doc, 'Required Metrics:')
    add_bullet(doc, 'Total Responses Analyzed')
    add_bullet(doc, 'Number of Codes Discovered')
    add_bullet(doc, 'Active Codes (codes with at least one assignment)')
    add_bullet(doc, 'Coverage Percentage (responses with at least one code)')
    add_bullet(doc, 'Average Codes per Response')
    add_bullet(doc, 'Average Confidence Score')
    add_bullet(doc, 'Confidence Score Range (min-max)')
    add_bullet(doc, 'Silhouette Score (if applicable)')
    add_bullet(doc, 'Execution Time')

    doc.add_paragraph()

    add_heading(doc, '3.3.5 Co-occurrence Tables', level=3)
    add_paragraph(doc, 'Purpose: Display pairs of codes that frequently appear together.')
    add_paragraph(doc, 'Required Columns:')
    add_bullet(doc, 'Code 1 (ID and Label)')
    add_bullet(doc, 'Code 2 (ID and Label)')
    add_bullet(doc, 'Co-occurrence Count')
    add_bullet(doc, 'Percentage of Total Responses')
    add_bullet(doc, 'Strength Indicator (e.g., Lift or Confidence)')

    doc.add_paragraph()

    add_heading(doc, '3.4 Chart Reporting Standards', level=2)
    add_paragraph(doc, 'Every chart must include the following elements:')

    chart_headers = ['Element', 'Requirement', 'Example']
    chart_rows = [
        ['Title', 'Descriptive and specific', 'Code Frequency Distribution (Top 15 Codes)'],
        ['Axis Labels', 'Include units if applicable', 'Number of Responses (n=1,234)'],
        ['Legend', 'If multiple series present', 'Algorithm: TF-IDF, LDA, NMF'],
        ['Caption', 'Below chart with interpretation', 'Figure 1: The most frequent code...'],
        ['Source Note', 'Data source and date', 'Source: Survey Data, December 2024'],
        ['Sample Size', 'Include in title or note', '(n=1,234 responses)'],
    ]
    add_table_content(doc, chart_headers, chart_rows)

    doc.add_paragraph()

    add_heading(doc, '3.5 Interpretation Notes', level=2)
    add_paragraph(doc, 'Every visualization and table should be accompanied by interpretation notes:')
    add_bullet(doc, 'Key Findings: 2-3 bullet points highlighting main insights')
    add_bullet(doc, 'Context: Background information needed to understand results')
    add_bullet(doc, 'Limitations: Any caveats or constraints on interpretation')
    add_bullet(doc, 'Recommendations: Suggested next steps or actions based on findings')

    doc.add_paragraph()

    add_paragraph(doc, 'Example interpretation note format:')
    add_paragraph(doc,
        'Key Findings: (1) The dominant code "Customer Service Quality" appears in 45.2% '
        'of responses, indicating strong thematic focus. (2) Average confidence score of 0.76 '
        'suggests high reliability in code assignments. (3) 12.3% of responses remain uncoded, '
        'warranting manual review for edge cases or novel themes.',
        italic=True)

    doc.add_page_break()

    # ============================================================================
    # 4. ANALYSIS INFERENCES AS TEXTUAL SUMMARIES
    # ============================================================================
    add_heading(doc, '4. Analysis Inferences as Textual Summaries', level=1)

    add_paragraph(doc,
        'Textual summaries transform quantitative results into narrative insights that are '
        'accessible to stakeholders at all levels of technical expertise. This section establishes '
        'standards for creating executive summaries, analytical narratives, and interpretive reports.')

    doc.add_paragraph()

    add_heading(doc, '4.1 Executive Summary Standards', level=2)
    add_paragraph(doc, 'Executive summaries should be concise, focused, and actionable.')

    doc.add_paragraph()
    add_paragraph(doc, 'Structure:')
    add_numbered(doc, 'Overview: 1-2 sentences describing the analysis scope and objectives')
    add_numbered(doc, 'Key Findings: 3-5 bullet points with the most important discoveries')
    add_numbered(doc, 'Implications: 2-3 sentences on what these findings mean')
    add_numbered(doc, 'Recommendations: 2-4 actionable next steps')

    doc.add_paragraph()
    add_paragraph(doc, 'Length: Maximum 300-400 words (one page)')
    add_paragraph(doc, 'Tone: Professional, objective, free of technical jargon')
    add_paragraph(doc, 'Audience: Senior leadership, non-technical stakeholders')

    doc.add_paragraph()

    add_heading(doc, '4.2 Analytical Narrative Standards', level=2)
    add_paragraph(doc, 'Analytical narratives provide detailed interpretation of results.')

    doc.add_paragraph()
    add_paragraph(doc, 'Required Sections:')
    add_numbered(doc, 'Methodology Summary: Brief description of analytical approach (150-200 words)')
    add_numbered(doc, 'Results Overview: Quantitative summary of main findings (200-300 words)')
    add_numbered(doc, 'Theme Analysis: Detailed discussion of each major theme/code (300-500 words)')
    add_numbered(doc, 'Patterns and Relationships: Co-occurrence and correlation insights (200-300 words)')
    add_numbered(doc, 'Quality Assessment: Discussion of confidence, coverage, and reliability (150-200 words)')
    add_numbered(doc, 'Limitations: Constraints, edge cases, and areas for improvement (100-150 words)')
    add_numbered(doc, 'Conclusions: Overall interpretation and significance (150-200 words)')

    doc.add_paragraph()
    add_paragraph(doc, 'Total Length: 1,500-2,000 words (3-4 pages)')

    doc.add_paragraph()

    add_heading(doc, '4.3 Insight Generation Framework', level=2)
    add_paragraph(doc, 'Use this framework to systematically generate insights:')

    doc.add_paragraph()
    add_paragraph(doc, 'Step 1: Identify Patterns')
    add_bullet(doc, 'What are the most frequent codes?')
    add_bullet(doc, 'Which codes frequently co-occur?')
    add_bullet(doc, 'Are there unexpected absences or presences?')

    doc.add_paragraph()
    add_paragraph(doc, 'Step 2: Contextualize Findings')
    add_bullet(doc, 'How do these patterns compare to expectations?')
    add_bullet(doc, 'What external factors might explain these patterns?')
    add_bullet(doc, 'How do results vary across segments or demographics?')

    doc.add_paragraph()
    add_paragraph(doc, 'Step 3: Assess Implications')
    add_bullet(doc, 'What do these findings mean for stakeholders?')
    add_bullet(doc, 'What opportunities or risks do they reveal?')
    add_bullet(doc, 'What decisions might be informed by these results?')

    doc.add_paragraph()
    add_paragraph(doc, 'Step 4: Formulate Recommendations')
    add_bullet(doc, 'What actions should be taken based on findings?')
    add_bullet(doc, 'What additional research or analysis is needed?')
    add_bullet(doc, 'What are the priorities for implementation?')

    doc.add_paragraph()

    add_heading(doc, '4.4 Automated Insight Templates', level=2)
    add_paragraph(doc,
        'The framework includes automated templates for common insights. These templates '
        'should be used as starting points and customized based on specific analytical context.')

    doc.add_paragraph()

    add_heading(doc, '4.4.1 Dominant Theme Insight', level=3)
    add_paragraph(doc, 'Template:', italic=True)
    add_paragraph(doc,
        'The dominant theme "[CODE_LABEL]" appears in [PERCENTAGE]% of responses ([COUNT] total), '
        'with an average confidence score of [CONFIDENCE]. This indicates [INTERPRETATION]. '
        'Representative examples include: "[QUOTE_1]" and "[QUOTE_2]".',
        italic=True)

    doc.add_paragraph()

    add_heading(doc, '4.4.2 Coverage Assessment Insight', level=3)
    add_paragraph(doc, 'Template for High Coverage (>80%):', italic=True)
    add_paragraph(doc,
        'The analysis achieved [COVERAGE]% coverage, successfully coding [CODED_COUNT] out of '
        '[TOTAL_COUNT] responses. This high coverage indicates that the discovered codes effectively '
        'capture the thematic diversity of the dataset.',
        italic=True)

    doc.add_paragraph()
    add_paragraph(doc, 'Template for Low Coverage (<80%):', italic=True)
    add_paragraph(doc,
        'The analysis achieved [COVERAGE]% coverage, with [UNCODED_COUNT] responses ([PERCENTAGE]%) '
        'remaining uncoded. These uncoded responses may represent novel themes not captured by current '
        'codes, measurement errors, or require manual review. Consider adjusting the confidence threshold '
        'or exploring additional codes.',
        italic=True)

    doc.add_paragraph()

    add_heading(doc, '4.4.3 Multi-Coding Insight', level=3)
    add_paragraph(doc, 'Template:', italic=True)
    add_paragraph(doc,
        '[MULTI_COUNT] responses ([PERCENTAGE]%) received multiple code assignments, averaging '
        '[AVG_CODES] codes per response. This indicates [INTERPRETATION: e.g., "nuanced perspectives", '
        '"complex narratives", "thematic overlap"]. The most common code pair is [CODE_1] and [CODE_2], '
        'co-occurring in [CO_COUNT] responses.',
        italic=True)

    doc.add_paragraph()

    add_heading(doc, '4.4.4 Quality Confidence Insight', level=3)
    add_paragraph(doc, 'Template:', italic=True)
    add_paragraph(doc,
        'Code assignments demonstrate [QUALITY_LEVEL: High/Moderate/Low] confidence with an average '
        'score of [AVG_CONFIDENCE]. This suggests that the ML algorithm is [INTERPRETATION: "highly '
        'certain", "moderately certain", "uncertain"] about its classifications. [RECOMMENDATION: '
        'Based on confidence level].',
        italic=True)

    doc.add_paragraph()

    add_heading(doc, '4.5 Narrative Best Practices', level=2)
    add_bullet(doc, 'Use Active Voice: "The analysis revealed..." rather than "It was revealed..."')
    add_bullet(doc, 'Include Specific Numbers: Provide exact counts and percentages, not vague terms')
    add_bullet(doc, 'Provide Context: Always explain what numbers mean in practical terms')
    add_bullet(doc, 'Use Comparative Language: Compare to benchmarks, expectations, or prior results')
    add_bullet(doc, 'Incorporate Examples: Include representative quotes to illustrate themes')
    add_bullet(doc, 'Maintain Objectivity: Distinguish between findings and interpretations')
    add_bullet(doc, 'Highlight Actionability: Focus on insights that can inform decisions')

    doc.add_page_break()

    # ============================================================================
    # 5. DASHBOARD AND UX INTEGRATION RECOMMENDATIONS
    # ============================================================================
    add_heading(doc, '5. Platform Dashboard and UX Integration Recommendations', level=1)

    add_paragraph(doc,
        'This section provides guidelines for integrating analytical outputs into interactive '
        'dashboards and user interfaces. Effective integration ensures that insights are accessible, '
        'actionable, and aligned with user workflows.')

    doc.add_paragraph()

    add_heading(doc, '5.1 Dashboard Architecture Principles', level=2)

    add_bullet(doc, 'Progressive Disclosure: Start with high-level summaries, allow drill-down to details')
    add_bullet(doc, 'User-Centric Design: Organize information around user tasks and questions')
    add_bullet(doc, 'Responsive Layout: Ensure functionality across desktop, tablet, and mobile devices')
    add_bullet(doc, 'Performance Optimization: Implement lazy loading and caching for large datasets')
    add_bullet(doc, 'Consistent Navigation: Maintain predictable patterns across all dashboard pages')

    doc.add_paragraph()

    add_heading(doc, '5.2 Recommended Dashboard Layout', level=2)

    doc.add_paragraph()
    add_heading(doc, '5.2.1 Overview Page', level=3)
    add_paragraph(doc, 'Purpose: Provide at-a-glance summary of key metrics.')
    add_paragraph(doc, 'Components:')
    add_numbered(doc, 'Header: Project/Analysis title, date range, sample size')
    add_numbered(doc, 'Key Metrics Cards: 4-6 cards showing total responses, codes discovered, coverage %, avg confidence')
    add_numbered(doc, 'Top Insights Panel: Auto-generated bullet points with key findings')
    add_numbered(doc, 'Primary Visualization: Code frequency bar chart (top 10-15 codes)')
    add_numbered(doc, 'Quick Actions: Export results, view full report, configure analysis')

    doc.add_paragraph()

    add_heading(doc, '5.2.2 Detailed Analysis Page', level=3)
    add_paragraph(doc, 'Purpose: Enable deep exploration of codes and themes.')
    add_paragraph(doc, 'Components:')
    add_numbered(doc, 'Interactive Codebook Table: Sortable, filterable, searchable')
    add_numbered(doc, 'Code Detail Panel: Click on code to view keywords, examples, statistics')
    add_numbered(doc, 'Visualization Tabs: Multiple chart types (frequency, scatter, histogram)')
    add_numbered(doc, 'Filter Controls: By confidence threshold, code frequency, date range')
    add_numbered(doc, 'Comparison Tools: Compare across segments, time periods, or subgroups')

    doc.add_paragraph()

    add_heading(doc, '5.2.3 Relationships and Patterns Page', level=3)
    add_paragraph(doc, 'Purpose: Explore co-occurrence and thematic relationships.')
    add_paragraph(doc, 'Components:')
    add_numbered(doc, 'Co-occurrence Heatmap: Interactive matrix with drill-down')
    add_numbered(doc, 'Network Diagram: Visualize code connections and clusters')
    add_numbered(doc, 'Top Pairs Table: Most frequent code combinations')
    add_numbered(doc, 'Relationship Insights: Auto-generated narrative on patterns')

    doc.add_paragraph()

    add_heading(doc, '5.2.4 Quality and Validation Page', level=3)
    add_paragraph(doc, 'Purpose: Assess reliability and identify areas for improvement.')
    add_paragraph(doc, 'Components:')
    add_numbered(doc, 'Quality Metrics Table: Coverage, confidence, silhouette score, etc.')
    add_numbered(doc, 'Confidence Distribution Chart: Histogram or violin plot')
    add_numbered(doc, 'Uncoded Responses Table: List of responses not assigned codes')
    add_numbered(doc, 'Validation Recommendations: Suggestions for improving results')

    doc.add_paragraph()

    add_heading(doc, '5.2.5 Export and Reporting Page', level=3)
    add_paragraph(doc, 'Purpose: Generate custom reports and export data.')
    add_paragraph(doc, 'Components:')
    add_numbered(doc, 'Format Selection: CSV, Excel, JSON, PDF report')
    add_numbered(doc, 'Content Customization: Select which outputs to include')
    add_numbered(doc, 'Template Selection: Executive summary, detailed report, technical documentation')
    add_numbered(doc, 'Preview Panel: Show preview of generated content')
    add_numbered(doc, 'Download Button: Initiate export with progress indicator')

    doc.add_paragraph()

    add_heading(doc, '5.3 Interactive Features', level=2)

    add_heading(doc, '5.3.1 Filtering and Search', level=3)
    add_bullet(doc, 'Global Search: Search across codes, keywords, and response text')
    add_bullet(doc, 'Multi-Criteria Filters: Combine filters for confidence, frequency, and segments')
    add_bullet(doc, 'Saved Filters: Allow users to save and reapply filter configurations')
    add_bullet(doc, 'Filter Persistence: Maintain filters across page navigation')

    doc.add_paragraph()

    add_heading(doc, '5.3.2 Drill-Down and Detail Views', level=3)
    add_bullet(doc, 'Click-to-Explore: Click on chart elements to see underlying data')
    add_bullet(doc, 'Modal Windows: Display detailed information without navigating away')
    add_bullet(doc, 'Breadcrumb Navigation: Show current location in navigation hierarchy')
    add_bullet(doc, 'Back Navigation: Easy return to previous views')

    doc.add_paragraph()

    add_heading(doc, '5.3.3 Customization Options', level=3)
    add_bullet(doc, 'Chart Type Selection: Allow users to toggle between chart types')
    add_bullet(doc, 'Color Theme: Light/dark mode or custom color schemes')
    add_bullet(doc, 'Layout Preferences: Adjust panel sizes and arrangement')
    add_bullet(doc, 'Saved Views: Save custom dashboard configurations')

    doc.add_paragraph()

    add_heading(doc, '5.4 User Experience Best Practices', level=2)

    add_heading(doc, '5.4.1 Loading and Performance', level=3)
    add_bullet(doc, 'Loading Indicators: Show progress bars or spinners during data loading')
    add_bullet(doc, 'Lazy Loading: Load visualizations only when visible')
    add_bullet(doc, 'Caching: Cache results to improve response times')
    add_bullet(doc, 'Pagination: Limit table rows to 50-100 per page')
    add_bullet(doc, 'Asynchronous Updates: Allow users to continue working during long operations')

    doc.add_paragraph()

    add_heading(doc, '5.4.2 Error Handling and Feedback', level=3)
    add_bullet(doc, 'Clear Error Messages: Explain what went wrong and how to fix it')
    add_bullet(doc, 'Validation Feedback: Provide immediate feedback on user inputs')
    add_bullet(doc, 'Success Confirmations: Confirm when actions complete successfully')
    add_bullet(doc, 'Warnings: Alert users to potential issues or data quality concerns')

    doc.add_paragraph()

    add_heading(doc, '5.4.3 Help and Documentation', level=3)
    add_bullet(doc, 'Contextual Help: Provide tooltips and info icons for complex features')
    add_bullet(doc, 'Onboarding Tutorial: Guide new users through key features')
    add_bullet(doc, 'Documentation Links: Link to relevant sections of user documentation')
    add_bullet(doc, 'FAQ Section: Address common questions and issues')

    doc.add_paragraph()

    add_heading(doc, '5.5 Mobile and Responsive Design', level=2)
    add_paragraph(doc, 'Guidelines for mobile and tablet devices:')
    add_bullet(doc, 'Responsive Grid: Use flexible layouts that adapt to screen size')
    add_bullet(doc, 'Touch-Friendly: Ensure buttons and interactive elements are at least 44x44 pixels')
    add_bullet(doc, 'Simplified Mobile View: Prioritize key metrics and hide secondary information')
    add_bullet(doc, 'Swipe Navigation: Support swipe gestures for tab and page navigation')
    add_bullet(doc, 'Mobile-Optimized Charts: Use simpler chart types and larger labels on mobile')

    doc.add_paragraph()

    add_heading(doc, '5.6 Integration with Existing Systems', level=2)

    add_heading(doc, '5.6.1 API Endpoints', level=3)
    add_paragraph(doc, 'Recommended REST API endpoints for integration:')
    add_bullet(doc, 'GET /api/analysis/{id}/summary - Retrieve executive summary')
    add_bullet(doc, 'GET /api/analysis/{id}/codes - Get codebook with all codes')
    add_bullet(doc, 'GET /api/analysis/{id}/assignments - Get response-level assignments')
    add_bullet(doc, 'GET /api/analysis/{id}/frequency - Get frequency table')
    add_bullet(doc, 'GET /api/analysis/{id}/cooccurrence - Get co-occurrence matrix')
    add_bullet(doc, 'GET /api/analysis/{id}/metrics - Get quality metrics')
    add_bullet(doc, 'POST /api/analysis - Create new analysis')
    add_bullet(doc, 'GET /api/analysis/{id}/export - Export results in specified format')

    doc.add_paragraph()

    add_heading(doc, '5.6.2 Embedding and Widgets', level=3)
    add_bullet(doc, 'Embeddable Visualizations: Provide iframe or JavaScript widgets for embedding')
    add_bullet(doc, 'OAuth Authentication: Secure API access with standard OAuth protocols')
    add_bullet(doc, 'Webhook Notifications: Notify external systems when analysis completes')
    add_bullet(doc, 'Data Connectors: Support connections to common data platforms (Tableau, PowerBI)')

    doc.add_paragraph()

    add_heading(doc, '5.7 Accessibility in Dashboards', level=2)
    add_bullet(doc, 'WCAG 2.1 AA Compliance: Meet Web Content Accessibility Guidelines')
    add_bullet(doc, 'Keyboard Navigation: All features accessible via keyboard')
    add_bullet(doc, 'Screen Reader Support: Proper ARIA labels and semantic HTML')
    add_bullet(doc, 'Alternative Data Tables: Provide tabular alternatives to visualizations')
    add_bullet(doc, 'High Contrast Mode: Support high contrast themes for visual impairment')
    add_bullet(doc, 'Text Sizing: Allow text scaling up to 200% without breaking layout')

    doc.add_page_break()

    # ============================================================================
    # 6. CONCLUSION
    # ============================================================================
    add_heading(doc, '6. Conclusion and Implementation Guidance', level=1)

    add_paragraph(doc,
        'These reporting and visualization standards establish a comprehensive framework for '
        'presenting qualitative coding analysis results with clarity, consistency, and professional '
        'quality. Adherence to these standards ensures that insights are accessible to diverse audiences, '
        'from technical analysts to executive stakeholders.')

    doc.add_paragraph()

    add_heading(doc, '6.1 Implementation Checklist', level=2)
    add_paragraph(doc, 'When implementing these standards, ensure:')

    add_bullet(doc, 'All visualizations use the approved color palette and typography')
    add_bullet(doc, 'Every table and chart includes required metadata (title, labels, source, sample size)')
    add_bullet(doc, 'Interpretation notes accompany all analytical outputs')
    add_bullet(doc, 'Textual summaries follow established templates and best practices')
    add_bullet(doc, 'Dashboard layouts adhere to progressive disclosure principles')
    add_bullet(doc, 'Interactive features enhance rather than complicate user experience')
    add_bullet(doc, 'Accessibility guidelines are met for all user interfaces')
    add_bullet(doc, 'Export and API functionality supports integration needs')

    doc.add_paragraph()

    add_heading(doc, '6.2 Continuous Improvement', level=2)
    add_paragraph(doc,
        'These standards should be reviewed and updated periodically based on:')
    add_bullet(doc, 'User feedback and usability testing results')
    add_bullet(doc, 'Emerging best practices in data visualization and UX design')
    add_bullet(doc, 'New analytical techniques and capabilities')
    add_bullet(doc, 'Evolving stakeholder needs and use cases')

    doc.add_paragraph()

    add_heading(doc, '6.3 Support and Resources', level=2)
    add_paragraph(doc, 'For questions or suggestions regarding these standards, consult:')
    add_bullet(doc, 'Project documentation and README files')
    add_bullet(doc, 'Example notebooks and code implementations')
    add_bullet(doc, 'Streamlit application for reference implementation')
    add_bullet(doc, 'Development team or project maintainers')

    doc.add_paragraph()
    doc.add_paragraph()

    add_paragraph(doc, '--- End of Document ---', bold=True)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save document
    output_path = '/home/user/JC-OE-Coding/documentation/Reporting_Visualization_Standards.docx'
    doc.save(output_path)
    print(f"Document created successfully: {output_path}")
    return output_path


if __name__ == '__main__':
    create_document()
