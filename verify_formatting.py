"""
Verify formatting applied to Benchmark_Standards.docx document.
"""

from docx import Document
from docx.shared import Pt, Inches


def verify_formatting(doc_path):
    """
    Verify the formatting applied to the document.

    Args:
        doc_path: Path to the Word document

    Returns:
        Dictionary with verification results
    """
    doc = Document(doc_path)

    print("=" * 70)
    print("DOCUMENT FORMATTING VERIFICATION")
    print("=" * 70)

    # Check margins
    print("\n1. MARGINS:")
    section = doc.sections[0]
    print(f"   Top Margin: {section.top_margin.inches:.2f} inches (Expected: 0.75)")
    print(f"   Right Margin: {section.right_margin.inches:.2f} inches (Expected: 0.75)")
    print(f"   Bottom Margin: {section.bottom_margin.inches:.2f} inches (Expected: 0.75)")
    print(f"   Left Margin: {section.left_margin.inches:.2f} inches (Expected: 0.75)")

    # Check default style
    print("\n2. DEFAULT STYLE (Normal):")
    style = doc.styles['Normal']
    print(f"   Font Name: {style.font.name} (Expected: Times New Roman)")
    print(f"   Font Size: {style.font.size.pt if style.font.size else 'Not set'}pt (Expected: 12pt)")

    # Sample paragraphs
    print("\n3. PARAGRAPH SAMPLES:")

    title_found = False
    h1_found = False
    h2_found = False
    h3_found = False
    body_found = False
    list_found = False

    for idx, para in enumerate(doc.paragraphs):
        if not title_found and (idx == 0 or (len(para.runs) > 0 and para.runs[0].font.size and para.runs[0].font.size >= Pt(20))):
            print("\n   TITLE (First paragraph or large font):")
            if para.runs:
                print(f"      Text: {para.text[:50]}...")
                print(f"      Font: {para.runs[0].font.name}")
                print(f"      Size: {para.runs[0].font.size.pt if para.runs[0].font.size else 'Not set'}pt (Expected: 28pt)")
                print(f"      Space Before: {para.paragraph_format.space_before.pt if para.paragraph_format.space_before else 0:.1f}pt (Expected: 12pt)")
                print(f"      Space After: {para.paragraph_format.space_after.pt if para.paragraph_format.space_after else 0:.1f}pt (Expected: 12pt)")
            title_found = True

        if not h1_found and 'Heading 1' in para.style.name:
            print("\n   HEADING 1:")
            print(f"      Text: {para.text}")
            if para.runs:
                print(f"      Font: {para.runs[0].font.name}")
                print(f"      Size: {para.runs[0].font.size.pt if para.runs[0].font.size else 'Not set'}pt (Expected: 16pt)")
                print(f"      Bold: {para.runs[0].font.bold}")
                print(f"      Space Before: {para.paragraph_format.space_before.pt if para.paragraph_format.space_before else 0:.1f}pt (Expected: 12pt)")
                print(f"      Space After: {para.paragraph_format.space_after.pt if para.paragraph_format.space_after else 0:.1f}pt (Expected: 6pt)")
            h1_found = True

        if not h2_found and 'Heading 2' in para.style.name:
            print("\n   HEADING 2:")
            print(f"      Text: {para.text}")
            if para.runs:
                print(f"      Font: {para.runs[0].font.name}")
                print(f"      Size: {para.runs[0].font.size.pt if para.runs[0].font.size else 'Not set'}pt (Expected: 14pt)")
                print(f"      Bold: {para.runs[0].font.bold}")
                print(f"      Space Before: {para.paragraph_format.space_before.pt if para.paragraph_format.space_before else 0:.1f}pt (Expected: 9pt)")
                print(f"      Space After: {para.paragraph_format.space_after.pt if para.paragraph_format.space_after else 0:.1f}pt (Expected: 5pt)")
            h2_found = True

        if not h3_found and 'Heading 3' in para.style.name:
            print("\n   HEADING 3:")
            print(f"      Text: {para.text}")
            if para.runs:
                print(f"      Font: {para.runs[0].font.name}")
                print(f"      Size: {para.runs[0].font.size.pt if para.runs[0].font.size else 'Not set'}pt (Expected: 13pt)")
                print(f"      Bold: {para.runs[0].font.bold}")
                print(f"      Space Before: {para.paragraph_format.space_before.pt if para.paragraph_format.space_before else 0:.1f}pt (Expected: 7pt)")
                print(f"      Space After: {para.paragraph_format.space_after.pt if para.paragraph_format.space_after else 0:.1f}pt (Expected: 4pt)")
            h3_found = True

        if not body_found and not para.style.name.startswith('Heading') and not para.style.name.startswith('List') and para.text.strip():
            print("\n   BODY TEXT:")
            print(f"      Text: {para.text[:50]}...")
            if para.runs:
                print(f"      Font: {para.runs[0].font.name}")
                print(f"      Size: {para.runs[0].font.size.pt if para.runs[0].font.size else 'Not set'}pt (Expected: 12pt)")
                print(f"      Line Spacing: {para.paragraph_format.line_spacing.pt if para.paragraph_format.line_spacing else 'Not set'}pt (Expected: 18pt for 1.5 spacing)")
            body_found = True

        if not list_found and para.style.name.startswith('List'):
            print("\n   LIST ITEM:")
            print(f"      Text: {para.text[:50]}...")
            if para.runs:
                print(f"      Font: {para.runs[0].font.name}")
                print(f"      Size: {para.runs[0].font.size.pt if para.runs[0].font.size else 'Not set'}pt (Expected: 12pt)")
                print(f"      Left Indent: {para.paragraph_format.left_indent.inches if para.paragraph_format.left_indent else 0:.2f} inches (Expected: 0.50)")
                print(f"      Hanging Indent: {para.paragraph_format.first_line_indent.inches if para.paragraph_format.first_line_indent else 0:.2f} inches (Expected: -0.25)")
            list_found = True

        if all([title_found, h1_found, h2_found, body_found, list_found]):
            break

    # Check tables
    print("\n4. TABLES:")
    print(f"   Total Tables: {len(doc.tables)}")

    if doc.tables:
        for table_idx, table in enumerate(doc.tables[:2]):  # Show first 2 tables
            print(f"\n   Table {table_idx + 1}:")
            print(f"      Rows: {len(table.rows)}")
            print(f"      Columns: {len(table.columns)}")

            # Check first cell formatting
            if table.rows:
                first_cell = table.rows[0].cells[0]
                print(f"      Header Cell Font: {first_cell.paragraphs[0].runs[0].font.name if first_cell.paragraphs and first_cell.paragraphs[0].runs else 'N/A'}")
                print(f"      Header Cell Bold: {first_cell.paragraphs[0].runs[0].font.bold if first_cell.paragraphs and first_cell.paragraphs[0].runs else 'N/A'}")
                print(f"      Header Alignment: {first_cell.paragraphs[0].paragraph_format.alignment if first_cell.paragraphs else 'N/A'}")

    print("\n" + "=" * 70)
    print("VERIFICATION COMPLETE")
    print("=" * 70)


def main():
    """Main function."""
    doc_path = '/home/user/JC-OE-Coding/documentation/Benchmark_Standards.docx'
    verify_formatting(doc_path)


if __name__ == '__main__':
    main()
