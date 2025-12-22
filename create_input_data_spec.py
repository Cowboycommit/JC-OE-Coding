"""
Script to create Input Data Specification Word document.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_input_data_specification():
    """Create a comprehensive Input Data Specification document."""

    doc = Document()

    # Set default font to Times New Roman
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Title
    title = doc.add_heading('Input Data Specification', 0)
    title_format = title.runs[0].font
    title_format.name = 'Times New Roman'
    title_format.size = Pt(16)
    title_format.bold = True

    # Subtitle
    subtitle = doc.add_paragraph('Open-Ended Coding Analysis Framework')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0].font
    subtitle_format.name = 'Times New Roman'
    subtitle_format.size = Pt(14)
    subtitle_format.italic = True

    doc.add_paragraph()

    # 1. Introduction
    heading1 = doc.add_heading('1. Introduction', 1)
    heading1.runs[0].font.name = 'Times New Roman'

    intro_text = (
        "This document specifies the data requirements, formatting standards, and preparation "
        "guidelines for the Open-Ended Coding Analysis Framework. The framework is designed to "
        "analyze qualitative data from surveys, interviews, and other open-ended response formats "
        "through systematic coding, theme identification, and hierarchical categorization."
    )
    p = doc.add_paragraph(intro_text)
    p.runs[0].font.name = 'Times New Roman'

    # 2. Dataset Requirements
    heading2 = doc.add_heading('2. Dataset Requirements', 1)
    heading2.runs[0].font.name = 'Times New Roman'

    # 2.1 Required Variables
    heading2_1 = doc.add_heading('2.1 Required Variables', 2)
    heading2_1.runs[0].font.name = 'Times New Roman'

    p = doc.add_paragraph('At minimum, your dataset must include:')
    p.runs[0].font.name = 'Times New Roman'

    required_vars = [
        ('response', 'Text', 'The open-ended text response from participants', 'Required'),
    ]

    for var_name, var_type, description, status in required_vars:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{var_name}').bold = True
        p.add_run(f' ({var_type}): {description} - ')
        p.add_run(status).italic = True
        for run in p.runs:
            run.font.name = 'Times New Roman'

    # 2.2 Optional Variables
    heading2_2 = doc.add_heading('2.2 Optional Variables', 2)
    heading2_2.runs[0].font.name = 'Times New Roman'

    p = doc.add_paragraph('The following variables are optional but recommended for enhanced analysis:')
    p.runs[0].font.name = 'Times New Roman'

    optional_vars = [
        ('id', 'Integer/String', 'Unique identifier for each response', 'Recommended'),
        ('respondent_id', 'String', 'Unique identifier for each participant (allows linking multiple responses)', 'Recommended'),
        ('timestamp', 'Date/DateTime', 'Date and time when the response was submitted', 'Optional'),
        ('topic', 'String', 'Pre-categorized topic or theme tag', 'Optional'),
        ('demographic_field', 'String/Categorical', 'Any demographic or segmentation variables (e.g., age, gender, region)', 'Optional'),
        ('weight', 'Float', 'Statistical weight for weighted analysis', 'Optional'),
    ]

    for var_name, var_type, description, status in optional_vars:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{var_name}').bold = True
        p.add_run(f' ({var_type}): {description} - ')
        p.add_run(status).italic = True
        for run in p.runs:
            run.font.name = 'Times New Roman'

    # 2.3 File Formats
    heading2_3 = doc.add_heading('2.3 Supported File Formats', 2)
    heading2_3.runs[0].font.name = 'Times New Roman'

    p = doc.add_paragraph('The framework supports the following input data formats:')
    p.runs[0].font.name = 'Times New Roman'

    formats = [
        ('CSV (Comma-Separated Values)', 'Preferred format. UTF-8 encoding recommended. File extension: .csv'),
        ('Excel', 'Microsoft Excel format. Supports .xlsx and .xls formats. Data should be in the first sheet or specify sheet name.'),
        ('JSON', 'JavaScript Object Notation. Supports both standard JSON and JSON Lines format. File extension: .json or .jsonl'),
        ('SQLite', 'SQLite database. Provide database file path and SQL query to extract data.'),
        ('PostgreSQL', 'PostgreSQL database. Provide connection string and SQL query to extract data.'),
    ]

    for format_name, description in formats:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(format_name).bold = True
        p.add_run(f': {description}')
        for run in p.runs:
            run.font.name = 'Times New Roman'

    # 2.4 Data Types
    heading2_4 = doc.add_heading('2.4 Data Types and Specifications', 2)
    heading2_4.runs[0].font.name = 'Times New Roman'

    data_types = [
        ('Text/String', 'Open-ended responses, identifiers, categorical variables. Encoding: UTF-8. No character limit, but responses longer than 10,000 characters may impact processing time.'),
        ('Integer', 'Numeric identifiers, counts. Range: Any valid integer.'),
        ('Float', 'Weights, scores, continuous variables. Format: Decimal notation (e.g., 1.5, 0.75).'),
        ('Date', 'Date values. Preferred format: YYYY-MM-DD (ISO 8601). Alternative formats: MM/DD/YYYY, DD/MM/YYYY (specify in documentation).'),
        ('DateTime', 'Date and time values. Preferred format: YYYY-MM-DD HH:MM:SS or ISO 8601 format with timezone.'),
        ('Categorical', 'Pre-defined categories or codes. Format: String values. Use consistent coding across all responses.'),
    ]

    for data_type, description in data_types:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(data_type).bold = True
        p.add_run(f': {description}')
        for run in p.runs:
            run.font.name = 'Times New Roman'

    # 3. Data Formatting Rules
    heading3 = doc.add_heading('3. Data Formatting Rules', 1)
    heading3.runs[0].font.name = 'Times New Roman'

    # 3.1 Variable Naming Conventions
    heading3_1 = doc.add_heading('3.1 Variable Naming Conventions', 2)
    heading3_1.runs[0].font.name = 'Times New Roman'

    p = doc.add_paragraph('Follow these naming conventions for variables (column names):')
    p.runs[0].font.name = 'Times New Roman'

    naming_rules = [
        'Use lowercase letters for variable names (e.g., response, respondent_id)',
        'Use underscores (_) to separate words (snake_case format)',
        'Avoid spaces, special characters, or punctuation in variable names',
        'Keep variable names concise but descriptive (e.g., timestamp not ts)',
        'Do not start variable names with numbers',
        'Use consistent naming across all variables',
        'Reserved names to avoid: id, response, codes, themes, categories (unless used for their intended purpose)',
    ]

    for rule in naming_rules:
        p = doc.add_paragraph(rule, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.add_run('Examples:').bold = True
    for run in p.runs:
        run.font.name = 'Times New Roman'

    examples = [
        ('Good: respondent_id, submission_date, age_group', True),
        ('Bad: Respondent ID, submission-date, ageGroup', False),
    ]

    for example, is_good in examples:
        p = doc.add_paragraph(example, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'

    # 3.2 Categorical Coding Standards
    heading3_2 = doc.add_heading('3.2 Categorical Coding Standards', 2)
    heading3_2.runs[0].font.name = 'Times New Roman'

    p = doc.add_paragraph('For categorical or pre-coded variables:')
    p.runs[0].font.name = 'Times New Roman'

    categorical_rules = [
        'Use consistent coding across all responses (e.g., always use "male" not "Male" or "M")',
        'Define a codebook for categorical variables with all possible values',
        'Prefer string labels over numeric codes for clarity (e.g., "agree" vs. "1")',
        'If using numeric codes, provide a separate codebook document',
        'Avoid mixing data types within a single variable',
        'Use standardized codes for common variables (e.g., ISO country codes)',
    ]

    for rule in categorical_rules:
        p = doc.add_paragraph(rule, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'

    # 3.3 Date Formats
    heading3_3 = doc.add_heading('3.3 Date and DateTime Formats', 2)
    heading3_3.runs[0].font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.add_run('Preferred format: ').font.name = 'Times New Roman'
    p.add_run('YYYY-MM-DD').bold = True
    p.runs[0].font.name = 'Times New Roman'
    p.runs[1].font.name = 'Times New Roman'

    p = doc.add_paragraph('For date and time values, use ISO 8601 format:')
    p.runs[0].font.name = 'Times New Roman'

    date_formats = [
        ('Date only', 'YYYY-MM-DD', '2024-04-15'),
        ('Date and time', 'YYYY-MM-DD HH:MM:SS', '2024-04-15 14:30:00'),
        ('Date and time with timezone', 'YYYY-MM-DDTHH:MM:SS+00:00', '2024-04-15T14:30:00+00:00'),
    ]

    for format_type, format_spec, example in date_formats:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{format_type}: ')
        p.add_run(format_spec).bold = True
        p.add_run(f' (Example: {example})')
        for run in p.runs:
            run.font.name = 'Times New Roman'

    p = doc.add_paragraph('Alternative acceptable formats (must be consistent throughout the dataset):')
    p.runs[0].font.name = 'Times New Roman'

    alt_formats = ['MM/DD/YYYY', 'DD/MM/YYYY', 'DD-MMM-YYYY']
    for fmt in alt_formats:
        p = doc.add_paragraph(fmt, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'

    # 3.4 Handling Missing Values
    heading3_4 = doc.add_heading('3.4 Handling Missing Values', 2)
    heading3_4.runs[0].font.name = 'Times New Roman'

    p = doc.add_paragraph('Missing data should be handled consistently:')
    p.runs[0].font.name = 'Times New Roman'

    missing_rules = [
        'Leave cells empty (blank) for missing values - this is the preferred method',
        'Alternatively, use standardized missing value codes: NA, N/A, or null',
        'Do not use: "missing", "unknown", "-", "0", or other non-standard codes',
        'Do not use numeric codes (e.g., -99, 999) for missing categorical data',
        'For required text responses, empty or missing responses will be flagged during validation',
        'Document any systematic patterns of missing data in your data preparation notes',
    ]

    for rule in missing_rules:
        p = doc.add_paragraph(rule, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'

    # 4. Data Quality Requirements
    heading4 = doc.add_heading('4. Data Quality Requirements', 1)
    heading4.runs[0].font.name = 'Times New Roman'

    quality_reqs = [
        ('Completeness', 'All required variables must be present. Text responses should not be empty unless genuinely non-responsive.'),
        ('Consistency', 'Use consistent formats, codes, and naming throughout the dataset.'),
        ('Encoding', 'Use UTF-8 encoding to ensure proper handling of special characters and international text.'),
        ('Duplicates', 'Remove duplicate responses unless intentional. Each response should have a unique identifier.'),
        ('Text Quality', 'Responses should contain meaningful text. Single-word or very short responses may limit analysis quality.'),
        ('File Size', 'CSV/Excel files should be under 500MB for optimal performance. For larger datasets, consider database formats.'),
    ]

    for req_name, description in quality_reqs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(req_name).bold = True
        p.add_run(f': {description}')
        for run in p.runs:
            run.font.name = 'Times New Roman'

    # 5. Data Structure Examples
    heading5 = doc.add_heading('5. Data Structure Examples', 1)
    heading5.runs[0].font.name = 'Times New Roman'

    # 5.1 Minimal Structure
    heading5_1 = doc.add_heading('5.1 Minimal Required Structure', 2)
    heading5_1.runs[0].font.name = 'Times New Roman'

    p = doc.add_paragraph('Simplest acceptable format (CSV example):')
    p.runs[0].font.name = 'Times New Roman'

    p = doc.add_paragraph('response')
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(10)

    p = doc.add_paragraph('"I love the flexibility of remote work"')
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(10)

    p = doc.add_paragraph('"Better work-life balance is crucial"')
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(10)

    # 5.2 Recommended Structure
    heading5_2 = doc.add_heading('5.2 Recommended Structure with Metadata', 2)
    heading5_2.runs[0].font.name = 'Times New Roman'

    p = doc.add_paragraph('Recommended format with optional fields (CSV example):')
    p.runs[0].font.name = 'Times New Roman'

    p = doc.add_paragraph('id,response,respondent_id,timestamp')
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(10)

    p = doc.add_paragraph('1,"I love the flexibility of remote work",R001,2024-04-01')
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(10)

    p = doc.add_paragraph('2,"Better work-life balance is crucial",R002,2024-04-02')
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(10)

    # 5.3 Enhanced Structure
    heading5_3 = doc.add_heading('5.3 Enhanced Structure with Segmentation Variables', 2)
    heading5_3.runs[0].font.name = 'Times New Roman'

    p = doc.add_paragraph('Complete format with demographic/segmentation variables (CSV example):')
    p.runs[0].font.name = 'Times New Roman'

    p = doc.add_paragraph('id,response,respondent_id,timestamp,topic,age_group,region')
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(10)

    p = doc.add_paragraph('1,"Sustainable fashion is essential",R001,2024-04-01,sustainability,25-34,northeast')
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(10)

    p = doc.add_paragraph('2,"Fast fashion creates waste",R002,2024-04-02,environment,35-44,west')
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(10)

    # 6. Templates and Schemas
    heading6 = doc.add_heading('6. Templates and Schemas', 1)
    heading6.runs[0].font.name = 'Times New Roman'

    p = doc.add_paragraph('The following templates are provided to guide data preparation:')
    p.runs[0].font.name = 'Times New Roman'

    templates = [
        ('Data Input Template (Excel)', 'A pre-formatted Excel template with proper column headers, data validation, and examples. See input_data_template.xlsx'),
        ('Sample Datasets', 'Example datasets demonstrating proper formatting. Available in the data/ directory.'),
        ('Validation Checklist', 'A checklist to verify your data meets all requirements before submission.'),
    ]

    for template_name, description in templates:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(template_name).bold = True
        p.add_run(f': {description}')
        for run in p.runs:
            run.font.name = 'Times New Roman'

    # 7. Data Validation
    heading7 = doc.add_heading('7. Data Validation Checklist', 1)
    heading7.runs[0].font.name = 'Times New Roman'

    p = doc.add_paragraph('Before submitting your data, verify the following:')
    p.runs[0].font.name = 'Times New Roman'

    validation_items = [
        'File format is one of the supported types (CSV, Excel, JSON, or database)',
        'Required "response" column is present and contains text data',
        'Variable names follow naming conventions (lowercase, underscores, no spaces)',
        'Date formats are consistent and use recommended format (YYYY-MM-DD)',
        'Missing values are handled consistently (blank cells or standardized codes)',
        'Categorical variables use consistent coding throughout',
        'File encoding is UTF-8 (for CSV and JSON files)',
        'No duplicate response IDs (if ID column is included)',
        'Text responses contain meaningful content (not empty or single characters)',
        'File size is manageable (under 500MB for CSV/Excel)',
        'All column headers are in the first row',
        'No merged cells or complex formatting (for Excel files)',
    ]

    for item in validation_items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'

    # 8. Common Issues and Solutions
    heading8 = doc.add_heading('8. Common Issues and Solutions', 1)
    heading8.runs[0].font.name = 'Times New Roman'

    issues = [
        ('Excel opens CSV with incorrect encoding', 'Save as CSV UTF-8 format in Excel. Alternatively, use a text editor to verify encoding.'),
        ('Special characters display incorrectly', 'Ensure file is saved with UTF-8 encoding. Avoid using system-specific encodings.'),
        ('Date values interpreted as text', 'Use consistent date format (YYYY-MM-DD) and ensure no leading/trailing spaces.'),
        ('Column names with spaces cause errors', 'Replace spaces with underscores. Use lowercase letters only.'),
        ('Missing value codes not recognized', 'Use blank cells or standard codes (NA, N/A). Do not use custom codes like "missing" or "-".'),
        ('Large file processing is slow', 'Split large files into smaller batches or use database formats for better performance.'),
    ]

    for issue, solution in issues:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'Issue: {issue}').bold = True
        p.runs[0].font.name = 'Times New Roman'
        p2 = doc.add_paragraph(f'Solution: {solution}', style='List Bullet 2')
        p2.runs[0].font.name = 'Times New Roman'

    # 9. Best Practices
    heading9 = doc.add_heading('9. Best Practices', 1)
    heading9.runs[0].font.name = 'Times New Roman'

    best_practices = [
        'Start with the provided template to ensure proper structure',
        'Test your data file with a small subset before preparing the full dataset',
        'Document any data cleaning or transformation steps you perform',
        'Keep a backup of your original raw data',
        'Include a data dictionary or codebook if using categorical variables',
        'Remove personally identifiable information (PII) before analysis',
        'Verify data quality using the validation checklist',
        'Use descriptive file names (e.g., survey_responses_2024_q1.csv)',
        'Consider pilot testing your data collection instrument to ensure quality responses',
    ]

    for practice in best_practices:
        p = doc.add_paragraph(practice, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'

    # 10. Support and Resources
    heading10 = doc.add_heading('10. Support and Resources', 1)
    heading10.runs[0].font.name = 'Times New Roman'

    p = doc.add_paragraph('For additional assistance:')
    p.runs[0].font.name = 'Times New Roman'

    resources = [
        ('Documentation', 'See the project README.md for detailed usage instructions'),
        ('Sample Data', 'Review example datasets in the data/ directory'),
        ('Template File', 'Use the Excel template (input_data_template.xlsx) as a starting point'),
        ('GitHub Repository', 'Visit https://github.com/Cowboycommit/JC-OE-Coding for updates and issues'),
    ]

    for resource_name, description in resources:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(resource_name).bold = True
        p.add_run(f': {description}')
        for run in p.runs:
            run.font.name = 'Times New Roman'

    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph('---')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.name = 'Times New Roman'

    footer_text = doc.add_paragraph('Open-Ended Coding Analysis Framework')
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text.runs[0].font.name = 'Times New Roman'
    footer_text.runs[0].font.size = Pt(10)
    footer_text.runs[0].font.italic = True

    version = doc.add_paragraph('Version 1.0 - December 2024')
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version.runs[0].font.name = 'Times New Roman'
    version.runs[0].font.size = Pt(10)

    # Save document
    output_path = '/home/user/JC-OE-Coding/documentation/Input_Data_Specification.docx'
    doc.save(output_path)
    print(f"Document created successfully: {output_path}")
    return output_path


if __name__ == '__main__':
    create_input_data_specification()
