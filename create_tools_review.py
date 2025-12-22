"""
Script to generate Open-Source Tools Review Word document
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_tools_review_document():
    """Create a comprehensive Open-Source Tools Review document"""

    # Create document
    doc = Document()

    # Set default font to Times New Roman
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Set narrow margins for more content on one page
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Title
    title = doc.add_heading('Open-Source Tools Review', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(16)
    title_run.font.bold = True

    # Subtitle
    subtitle = doc.add_paragraph('Python-Based Libraries for Qualitative Data Analysis')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.name = 'Times New Roman'
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.italic = True

    # Introduction
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        'This document provides an overview of reliable Python-based open-source libraries '
        'used in the Open-Ended Coding Analysis framework for qualitative data analysis, '
        'machine learning-based theme discovery, and interactive visualization.'
    )
    intro_run.font.name = 'Times New Roman'
    intro_run.font.size = Pt(11)

    # Section 1: Data Manipulation
    heading1 = doc.add_heading('1. Data Manipulation and Processing', level=2)
    heading1.runs[0].font.name = 'Times New Roman'
    heading1.runs[0].font.size = Pt(12)

    p1 = doc.add_paragraph()
    p1_run = p1.add_run(
        'Pandas (v2.0.0+): Industry-standard data manipulation library providing DataFrame structures '
        'and comprehensive data analysis tools. Selected for its robust CSV/Excel handling, SQL integration, '
        'and extensive data transformation capabilities. NumPy (v1.24.0+) provides foundational numerical '
        'computing support with efficient array operations and mathematical functions.'
    )
    p1_run.font.name = 'Times New Roman'
    p1_run.font.size = Pt(11)

    # Section 2: Machine Learning and NLP
    heading2 = doc.add_heading('2. Machine Learning and Natural Language Processing', level=2)
    heading2.runs[0].font.name = 'Times New Roman'
    heading2.runs[0].font.size = Pt(12)

    p2 = doc.add_paragraph()
    p2_run = p2.add_run(
        'Scikit-learn (v1.3.0+): Best choice for ML-based coding with comprehensive algorithms including '
        'TF-IDF vectorization, K-Means clustering, Latent Dirichlet Allocation (LDA), and Non-negative '
        'Matrix Factorization (NMF). Justification: Proven reliability, extensive documentation, excellent '
        'performance, and seamless integration with pandas. NLTK (v3.8.0+) provides text preprocessing, '
        'tokenization, and stopword removal capabilities essential for qualitative text analysis.'
    )
    p2_run.font.name = 'Times New Roman'
    p2_run.font.size = Pt(11)

    # Section 3: Visualization
    heading3 = doc.add_heading('3. Data Visualization', level=2)
    heading3.runs[0].font.name = 'Times New Roman'
    heading3.runs[0].font.size = Pt(12)

    p3 = doc.add_paragraph()
    p3_run = p3.add_run(
        'Plotly (v5.14.0+): Best choice for interactive visualizations with professional-quality output '
        'suitable for publication and stakeholder presentations. Justification: Superior interactivity, '
        'web-ready exports, and extensive chart types. Matplotlib (v3.7.0+) and Seaborn (v0.12.0+) provide '
        'statistical plotting and publication-quality static figures. NetworkX (v3.1+) enables code co-occurrence '
        'network analysis. WordCloud (v1.9.0+) generates thematic word clouds for qualitative insights.'
    )
    p3_run.font.name = 'Times New Roman'
    p3_run.font.size = Pt(11)

    # Section 4: Web Interface
    heading4 = doc.add_heading('4. Web Application Framework', level=2)
    heading4.runs[0].font.name = 'Times New Roman'
    heading4.runs[0].font.size = Pt(12)

    p4 = doc.add_paragraph()
    p4_run = p4.add_run(
        'Streamlit (v1.28.0+): Best choice for rapid development of data science web applications. '
        'Justification: Zero HTML/CSS/JavaScript required, native support for data science libraries, '
        'built-in caching mechanisms, and excellent performance for interactive dashboards. Enables '
        'non-programmers to perform complex ML-based coding analysis through intuitive drag-and-drop interfaces.'
    )
    p4_run.font.name = 'Times New Roman'
    p4_run.font.size = Pt(11)

    # Section 5: Database Connectivity
    heading5 = doc.add_heading('5. Database Integration', level=2)
    heading5.runs[0].font.name = 'Times New Roman'
    heading5.runs[0].font.size = Pt(12)

    p5 = doc.add_paragraph()
    p5_run = p5.add_run(
        'SQLAlchemy (v2.0.0+): Best choice for database abstraction with support for multiple database backends '
        '(SQLite, PostgreSQL, MySQL). Justification: Database-agnostic API, robust connection pooling, and '
        'excellent ORM capabilities. Psycopg2-binary (v2.9.0+) provides optimized PostgreSQL connectivity for '
        'large-scale qualitative datasets.'
    )
    p5_run.font.name = 'Times New Roman'
    p5_run.font.size = Pt(11)

    # Section 6: File Format Support
    heading6 = doc.add_heading('6. File Format Handling', level=2)
    heading6.runs[0].font.name = 'Times New Roman'
    heading6.runs[0].font.size = Pt(12)

    p6 = doc.add_paragraph()
    p6_run = p6.add_run(
        'OpenPyXL (v3.1.0+): Best choice for Excel file operations with support for .xlsx format, '
        'multiple worksheets, and formatting preservation. Justification: Comprehensive Excel feature support, '
        'active development, and seamless pandas integration. XLRD (v2.0.1+) provides legacy .xls format support. '
        'PyYAML (v6.0+) enables structured configuration management.'
    )
    p6_run.font.name = 'Times New Roman'
    p6_run.font.size = Pt(11)

    # Section 7: Development Tools
    heading7 = doc.add_heading('7. Development and Quality Assurance', level=2)
    heading7.runs[0].font.name = 'Times New Roman'
    heading7.runs[0].font.size = Pt(12)

    p7 = doc.add_paragraph()
    p7_run = p7.add_run(
        'Pytest (v7.3.0+): Best choice for testing framework with extensive plugin ecosystem, fixtures, '
        'and parametrization. Black (v23.3.0+) ensures consistent code formatting. Jupyter (v1.0.0+) and '
        'Notebook (v7.0.0+) provide interactive development environments essential for iterative qualitative '
        'analysis workflows.'
    )
    p7_run.font.name = 'Times New Roman'
    p7_run.font.size = Pt(11)

    # Conclusion
    conclusion_heading = doc.add_heading('Conclusion', level=2)
    conclusion_heading.runs[0].font.name = 'Times New Roman'
    conclusion_heading.runs[0].font.size = Pt(12)

    conclusion = doc.add_paragraph()
    conclusion_run = conclusion.add_run(
        'The selected libraries represent industry-standard, well-maintained open-source tools with proven '
        'reliability in production environments. This technology stack enables comprehensive qualitative data '
        'analysis from data ingestion through ML-based coding to interactive visualization and stakeholder reporting. '
        'The combination of scikit-learn for ML algorithms, Plotly for visualization, and Streamlit for web interfaces '
        'provides the optimal balance of functionality, performance, and usability for qualitative research applications.'
    )
    conclusion_run.font.name = 'Times New Roman'
    conclusion_run.font.size = Pt(11)

    # Save document
    output_path = '/home/user/JC-OE-Coding/documentation/Open-Source_Tools_Review.docx'
    doc.save(output_path)
    print(f"Document created successfully: {output_path}")

    return output_path

if __name__ == '__main__':
    create_tools_review_document()
