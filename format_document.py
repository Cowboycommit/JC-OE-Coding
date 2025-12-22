#!/usr/bin/env python3
"""
Format Word document according to specifications.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_margins(cell, top=100, bottom=100, left=180, right=180):
    """Set cell margins in twentieths of a point."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for margin_name, value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)

    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    for edge in ('top', 'left', 'bottom', 'right'):
        edge_element = OxmlElement(f'w:{edge}')
        edge_element.set(qn('w:val'), 'single')
        edge_element.set(qn('w:sz'), '8')  # 1pt = 8 eighths of a point
        edge_element.set(qn('w:space'), '0')
        edge_element.set(qn('w:color'), '000000')
        tcBorders.append(edge_element)

    tcPr.append(tcBorders)

def set_cell_background(cell, color):
    """Set cell background color."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    tcPr.append(shading)

def format_document(doc_path):
    """Apply formatting specifications to document."""
    doc = Document(doc_path)

    # Set page margins (0.75 inches = 1080 DXA)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Track statistics
    stats = {
        'paragraphs': 0,
        'headings': {'title': 0, 'h1': 0, 'h2': 0, 'h3': 0},
        'tables': 0,
        'body_text': 0,
        'lists': 0
    }

    # Format paragraphs
    for para in doc.paragraphs:
        stats['paragraphs'] += 1

        # Get paragraph style
        style_name = para.style.name if para.style else ''

        # Set font to Times New Roman for all runs
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)

            # Preserve bold and italic
            # Keep code formatting with Courier New
            if hasattr(run.font, 'name') and 'Courier' in str(run.font.name):
                run.font.name = 'Courier New'

        # Apply formatting based on style
        if 'Title' in style_name:
            stats['headings']['title'] += 1
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.size = Pt(28)  # 28pt
                run.font.name = 'Times New Roman'
                # Apply bold and italic to title
                run.font.bold = True
                run.font.italic = True
            # Title spacing
            para.paragraph_format.space_before = Pt(240)  # 240pt
            para.paragraph_format.space_after = Pt(240)   # 240pt
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        elif 'Heading 1' in style_name:
            stats['headings']['h1'] += 1
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.size = Pt(16)  # 16pt
                run.font.name = 'Times New Roman'
            # H1 spacing
            para.paragraph_format.space_before = Pt(240)  # 240pt
            para.paragraph_format.space_after = Pt(120)   # 120pt
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        elif 'Heading 2' in style_name:
            stats['headings']['h2'] += 1
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.size = Pt(14)  # 14pt
                run.font.name = 'Times New Roman'
            # H2 spacing
            para.paragraph_format.space_before = Pt(180)  # 180pt
            para.paragraph_format.space_after = Pt(100)   # 100pt
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        elif 'Heading 3' in style_name:
            stats['headings']['h3'] += 1
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.size = Pt(13)  # 13pt
                run.font.name = 'Times New Roman'
            # H3 spacing
            para.paragraph_format.space_before = Pt(140)  # 140pt
            para.paragraph_format.space_after = Pt(80)    # 80pt
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        elif 'TOC' in style_name or 'toc' in style_name.lower():
            # Table of Contents formatting
            if 'TOC 1' in style_name or 'toc 1' in style_name.lower():
                # Main level: Indent Left: 360pt, Hanging Indent: 360pt
                para.paragraph_format.left_indent = Pt(360)
                para.paragraph_format.first_line_indent = Pt(-360)
            else:
                # Sub level: Indent Left: 1080pt, Hanging Indent: 720pt
                para.paragraph_format.left_indent = Pt(1080)
                para.paragraph_format.first_line_indent = Pt(-720)

            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            for run in para.runs:
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'

        elif 'List' in style_name or para.text.strip().startswith('•'):
            stats['lists'] += 1
            # List formatting: Indent Left: 720pt, Hanging Indent: 360pt
            para.paragraph_format.left_indent = Pt(720)
            para.paragraph_format.first_line_indent = Pt(-360)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            for run in para.runs:
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'

        else:
            # Body text
            if para.text.strip():  # Only count non-empty paragraphs
                stats['body_text'] += 1
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                if 'Courier' not in str(run.font.name):  # Preserve code blocks
                    run.font.size = Pt(12)  # 12pt
                    run.font.name = 'Times New Roman'
            # Body text spacing
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # Format tables
    for table in doc.tables:
        stats['tables'] += 1

        # Calculate column widths (equal distribution)
        total_width = Inches(6.5)  # 8" - 1.5" margins = 6.5"
        num_cols = len(table.columns)
        col_width = total_width / num_cols

        for col in table.columns:
            col.width = int(col_width)

        # Format each row
        for row_idx, row in enumerate(table.rows):
            is_header = (row_idx == 0)

            for cell in row.cells:
                # Set cell borders
                set_cell_border(cell)

                # Set cell margins
                set_cell_margins(cell, top=100, bottom=100, left=180, right=180)

                # Set background color
                if is_header:
                    set_cell_background(cell, 'DDDDDD')
                else:
                    set_cell_background(cell, 'FFFFFF')

                # Format cell paragraphs
                for para in cell.paragraphs:
                    # Set alignment
                    if is_header:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    # Set line spacing (1.0 for tables)
                    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

                    # Format runs
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(0, 0, 0)

                        # Bold for headers
                        if is_header:
                            run.font.bold = True

    # Save the document
    doc.save(doc_path)

    return stats

if __name__ == '__main__':
    doc_path = '/home/user/JC-OE-Coding/documentation/Reporting_Visualization_Standards.docx'

    print(f"Formatting document: {doc_path}")
    print("-" * 60)

    stats = format_document(doc_path)

    print("\nFormatting Complete!")
    print("=" * 60)
    print(f"Total Paragraphs: {stats['paragraphs']}")
    print(f"  - Title: {stats['headings']['title']}")
    print(f"  - Heading 1: {stats['headings']['h1']}")
    print(f"  - Heading 2: {stats['headings']['h2']}")
    print(f"  - Heading 3: {stats['headings']['h3']}")
    print(f"  - Body Text: {stats['body_text']}")
    print(f"  - Lists: {stats['lists']}")
    print(f"Total Tables: {stats['tables']}")
    print("=" * 60)
    print("\nApplied Specifications:")
    print("  ✓ Font: Times New Roman (throughout entire document)")
    print("  ✓ Font Sizes: Title (28pt), H1 (16pt), H2 (14pt), H3 (13pt), Body (12pt)")
    print("  ✓ Line Spacing: Body (1.5), Tables (1.0), TOC (1.5)")
    print("  ✓ Margins: 0.75\" all sides (1080 DXA)")
    print("  ✓ Heading spacing: Title (240/240pt), H1 (240/120pt), H2 (180/100pt), H3 (140/80pt)")
    print("  ✓ Table formatting: borders (1pt black), backgrounds (grey/white), cell margins")
    print("  ✓ Text colors: Black (#000000)")
    print("  ✓ Alignment: Title/Headings (left), Table headers (center), Table content (left)")
    print("  ✓ List formatting: Bullet indents (720pt left, 360pt hanging)")
    print("  ✓ TOC formatting: Main (360pt left, 360pt hanging), Sub (1080pt left, 720pt hanging)")
